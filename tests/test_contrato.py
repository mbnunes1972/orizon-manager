import sys, os, json
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

import pytest
from unittest.mock import patch, MagicMock
from mod_contrato import calcular_hash_assinatura, montar_variaveis_contrato, gerar_pdf_contrato


@pytest.fixture(autouse=True)
def _isola_contratos_dir(tmp_path):
    """Isola CONTRATOS_DIR num diretório temporário para que a geração de contratos nos testes
    não escreva na pasta CONTRATOS do repositório (era a origem dos contrato_99/8888/9999.docx)."""
    import mod_contrato
    orig = mod_contrato.CONTRATOS_DIR
    mod_contrato.CONTRATOS_DIR = str(tmp_path / "contratos")
    os.makedirs(mod_contrato.CONTRATOS_DIR, exist_ok=True)
    yield
    mod_contrato.CONTRATOS_DIR = orig


def test_hash_assinatura_determinístico():
    h1 = calcular_hash_assinatura("João Silva", "123.456.789-00", 42, "2026-06-15T10:00:00")
    h2 = calcular_hash_assinatura("João Silva", "123.456.789-00", 42, "2026-06-15T10:00:00")
    assert h1 == h2


def test_hash_assinatura_muda_com_dados_diferentes():
    h1 = calcular_hash_assinatura("João Silva", "123.456.789-00", 42, "2026-06-15T10:00:00")
    h2 = calcular_hash_assinatura("Maria Silva", "123.456.789-00", 42, "2026-06-15T10:00:00")
    assert h1 != h2


def test_hash_assinatura_formato_sha256():
    h = calcular_hash_assinatura("João", "000.000.000-00", 1, "2026-01-01T00:00:00")
    assert len(h) == 64
    assert all(c in "0123456789abcdef" for c in h)


def test_montar_variaveis_contrato_campos_obrigatorios():
    variaveis = montar_variaveis_contrato(
        projeto={"nome_projeto": "Cozinha Silva", "criado_em": "2026-06-15", "consultor": "Pedro"},
        cliente={"nome": "Ana Silva", "cpf": "123.456.789-00", "telefone": "(11) 99999-9999",
                 "logradouro": "Rua A", "numero": "100", "bairro": "Centro", "cidade": "SP", "estado": "SP"},
        orcamento={"nome": "Orçamento 1", "valor_total": 48200.0, "forma_pagamento": "", "ambientes": []},
        endereco_instalacao="Rua B, 200 - Centro - SP",
        entrada_valor=5000.0,
        parcelas_descricao="11x",
        adendo="",
    )
    assert variaveis["cliente_nome"] == "Ana Silva"
    assert variaveis["cliente_cpf"]  == "123.456.789-00"
    assert variaveis["consultor_nome"] == "Pedro"
    assert variaveis["adendo"] == ""


def test_montar_variaveis_sem_adendo_retorna_string_vazia():
    variaveis = montar_variaveis_contrato(
        projeto={"nome_projeto": "P", "criado_em": "2026-01-01", "consultor": "X"},
        cliente={"nome": "C", "cpf": "", "telefone": "", "logradouro": "",
                 "numero": "", "bairro": "", "cidade": "", "estado": ""},
        orcamento={"nome": "O", "valor_total": 0.0, "forma_pagamento": "", "ambientes": []},
        endereco_instalacao="", entrada_valor=0.0, parcelas_descricao="", adendo=None,
    )
    assert variaveis["adendo"] == ""


def test_gerar_pdf_usa_modelo_e_chama_libreoffice(tmp_path):
    from mod_contrato import construir_contexto, _MODELO
    # Cria um modelo fake se não existir no ambiente de teste
    if not os.path.exists(_MODELO):
        pytest_mark = "skip"
        return  # pula se não há modelo no ambiente de CI
    ctx = construir_contexto(
        cliente={"nome": "Teste", "cpf": "000", "email": "", "telefone": "",
                 "logradouro": "", "numero": "", "complemento": "", "bairro": "",
                 "cidade": "", "cep": "", "estado": "",
                 "inst_mesmo_residencial": True,
                 "inst_logradouro": "", "inst_numero": "", "inst_complemento": "",
                 "inst_bairro": "", "inst_cidade": "", "inst_cep": "", "inst_uf": ""},
        usuario={"nome": "Consultor X", "telefone": "", "email": ""},
        forma_pagamento_json="",
    )
    with patch("mod_contrato.subprocess.run") as mock_run:
        mock_run.return_value = MagicMock(returncode=0)
        result = gerar_pdf_contrato(contrato_id=99, variaveis=ctx)
    assert "99" in result
    mock_run.assert_called_once()
    run_args = mock_run.call_args[0][0]
    assert "--convert-to" in run_args


def test_construir_contexto_aymore():
    from mod_contrato import construir_contexto
    cliente = {
        "nome": "João Silva", "cpf": "123.456.789-00",
        "email": "joao@test.com", "telefone": "12999990000",
        "logradouro": "Rua A", "numero": "10", "complemento": "",
        "bairro": "Centro", "cidade": "SJC", "cep": "12200-000", "estado": "SP",
        "inst_mesmo_residencial": True,
        "inst_logradouro": "", "inst_numero": "", "inst_complemento": "",
        "inst_bairro": "", "inst_cidade": "", "inst_cep": "", "inst_uf": "",
    }
    usuario = {"nome": "Pedro", "telefone": None, "email": "pedro@loja.com"}
    # Estrutura NOVA (_capturarPagamento): parcelas com valor numérico, total_cliente
    forma = json.dumps({
        "tipo": "aymore",
        "nome_forma": "Financiamento Aymoré",
        "entrada_valor": 5000.0,
        "entrada_forma": "Boleto",
        "entrada_data": "2026-07-15",
        "total_cliente": 3000.0,
        "texto_cartao": "",
        "parcelas": [
            {"num": 1, "data": "15/08/2026", "valor": 1000.0},
            {"num": 2, "data": "15/09/2026", "valor": 1000.0},
            {"num": 3, "data": "15/10/2026", "valor": 1000.0},
        ]
    })
    ctx = construir_contexto(cliente, usuario, forma, {"telefone": "(12) 3341-8777", "email": "sac@dalmobilesjc.com.br"})
    assert ctx["consultor_nome"] == "Pedro"
    assert ctx["consultor_tel"] == "(12) 3341-8777"   # fallback
    assert ctx["consultor_email"] == "pedro@loja.com"
    assert ctx["cliente_nome"] == "João Silva"
    assert ctx["inst_logradouro"] == "Rua A"           # mesmo endereço residencial
    pag = ctx["_pag"]
    assert pag["entrada_valor"] == "R$ 5.000,00"
    assert pag["valor_contrato"] == "R$ 3.000,00"
    assert pag["num_parcelas_int"] == 3
    assert pag["valores"][0] == "R$ 1.000,00"
    assert pag["datas"][0] == "15/08/2026"
    assert pag["datas"][1] == "15/09/2026"
    assert pag["datas"][3] == ""                       # parcelas além de 3 = ""
    assert "data_contrato" in ctx


def test_construir_contexto_cartao():
    from mod_contrato import construir_contexto
    cliente = {
        "nome": "Ana", "cpf": "", "email": "", "telefone": "",
        "logradouro": "", "numero": "", "complemento": "", "bairro": "",
        "cidade": "", "cep": "", "estado": "",
        "inst_mesmo_residencial": True,
        "inst_logradouro": "", "inst_numero": "", "inst_complemento": "",
        "inst_bairro": "", "inst_cidade": "", "inst_cep": "", "inst_uf": "",
    }
    usuario = {"nome": "Luiz", "telefone": "12988880000", "email": ""}
    forma = json.dumps({
        "tipo": "cartao", "nome_forma": "Cartão de Crédito",
        "entrada_valor": 0, "entrada_data": "", "parcelas": [],
        "texto_cartao": "12x R$ 10.000,00", "total_cliente": 120000,
    })
    ctx = construir_contexto(cliente, usuario, forma, {"telefone": "(12) 3341-8777", "email": "sac@dalmobilesjc.com.br"})
    assert ctx["consultor_tel"] == "12988880000"
    assert ctx["consultor_email"] == "sac@dalmobilesjc.com.br"   # fallback email
    pag = ctx["_pag"]
    assert pag["num_parcelas_int"] == 0
    assert pag["texto_cartao"] == "12x R$ 10.000,00"
    assert pag["valor_contrato"] == "R$ 120.000,00"
    assert pag["datas"] == [""] * 24
    assert pag["valores"] == [""] * 24


def test_construir_contexto_total_flex():
    from mod_contrato import construir_contexto
    cliente = {
        "nome": "Carlos", "cpf": "", "email": "", "telefone": "",
        "logradouro": "Av B", "numero": "5", "complemento": "", "bairro": "Vila",
        "cidade": "SP", "cep": "01000-000", "estado": "SP",
        "inst_mesmo_residencial": False,
        "inst_logradouro": "Rua C", "inst_numero": "20", "inst_complemento": "Ap 1",
        "inst_bairro": "Jardim", "inst_cidade": "SP", "inst_cep": "02000-000", "inst_uf": "SP",
    }
    usuario = {"nome": "Marcia", "telefone": "", "email": ""}
    forma = json.dumps({
        "tipo": "tf", "nome_forma": "Total Flex",
        "entrada_valor": 3000.0, "entrada_data": "01/07/2026",
        "total_cliente": 10000.0, "texto_cartao": "",
        "parcelas": [
            {"num": i, "data": f"10/{6+i:02d}/2026", "valor": 2000.0}
            for i in range(1, 6)
        ]
    })
    ctx = construir_contexto(cliente, usuario, forma, {"telefone": "(12) 3341-8777", "email": "sac@dalmobilesjc.com.br"})
    assert ctx["consultor_tel"] == "(12) 3341-8777"
    assert ctx["inst_logradouro"] == "Rua C"
    assert ctx["res_logradouro"] == "Av B"
    pag = ctx["_pag"]
    assert pag["num_parcelas_int"] == 5
    assert pag["datas"][0] == "10/07/2026"
    assert pag["datas"][4] == "10/11/2026"
    assert pag["datas"][5] == ""
    assert pag["valores"][0] == "R$ 2.000,00"


def _cliente_completo():
    """Cliente com todos os campos obrigatórios para gerar contrato."""
    return {
        "nome": "Ana Silva", "cpf": "123.456.789-00",
        "email": "ana@test.com", "telefone": "(12) 99999-0000",
        "logradouro": "Rua A", "numero": "100", "complemento": "",
        "bairro": "Centro", "cidade": "SJC", "cep": "12200-000", "estado": "SP",
        "inst_mesmo_residencial": True,
        "inst_logradouro": "", "inst_numero": "", "inst_complemento": "",
        "inst_bairro": "", "inst_cidade": "", "inst_cep": "", "inst_uf": "",
    }


def test_validar_cliente_completo_sem_faltas():
    from mod_contrato import validar_cliente_para_contrato
    assert validar_cliente_para_contrato(_cliente_completo()) == []


def test_validar_cliente_sem_endereco_residencial():
    from mod_contrato import validar_cliente_para_contrato
    c = _cliente_completo()
    for campo in ("logradouro", "numero", "bairro", "cidade", "cep", "estado"):
        c[campo] = ""
    faltando = validar_cliente_para_contrato(c)
    # Todos os 6 campos residenciais devem ser apontados como faltando
    assert len(faltando) == 6
    joined = " ".join(faltando).lower()
    for termo in ("logradouro", "número", "bairro", "cidade", "cep", "estado"):
        assert termo in joined


def test_validar_cliente_sem_contato():
    from mod_contrato import validar_cliente_para_contrato
    c = _cliente_completo()
    c["email"] = ""
    c["telefone"] = None
    faltando = validar_cliente_para_contrato(c)
    joined = " ".join(faltando).lower()
    assert "e-mail" in joined
    assert "telefone" in joined


def test_validar_inst_diferente_exige_endereco_instalacao():
    from mod_contrato import validar_cliente_para_contrato
    c = _cliente_completo()
    c["inst_mesmo_residencial"] = False
    # inst_* vazios → devem ser cobrados
    faltando = validar_cliente_para_contrato(c)
    joined = " ".join(faltando).lower()
    assert "instalação" in joined
    # Preenchendo inst_* → sem faltas
    c.update({"inst_logradouro": "Rua C", "inst_numero": "20", "inst_bairro": "Jardim",
              "inst_cidade": "SP", "inst_cep": "02000-000", "inst_uf": "SP"})
    assert validar_cliente_para_contrato(c) == []


def test_validar_inst_mesma_nao_exige_inst_fields():
    from mod_contrato import validar_cliente_para_contrato
    c = _cliente_completo()  # inst_mesmo_residencial=True, inst_* vazios
    assert validar_cliente_para_contrato(c) == []


def test_email_fallback_consultor():
    from mod_contrato import construir_contexto
    cliente = {"nome": "X", "cpf": "", "email": "", "telefone": "",
               "logradouro": "", "numero": "", "complemento": "", "bairro": "",
               "cidade": "", "cep": "", "estado": "",
               "inst_mesmo_residencial": True,
               "inst_logradouro": "", "inst_numero": "", "inst_complemento": "",
               "inst_bairro": "", "inst_cidade": "", "inst_cep": "", "inst_uf": ""}
    loja = {"telefone": "(12) 3341-8777", "email": "sac@dalmobilesjc.com.br"}
    ctx = construir_contexto(cliente, {"nome": "X", "telefone": "", "email": ""}, "", loja)
    assert ctx["consultor_email"] == "sac@dalmobilesjc.com.br"
    assert ctx["consultor_tel"]   == "(12) 3341-8777"


def test_preencher_signatario_e_testemunhas(tmp_path):
    import os
    from mod_contrato import preencher_contrato, _MODELO, construir_contexto
    if not os.path.exists(_MODELO):
        return
    from docx import Document
    loja = {"nome": "INSPIRIUM MOVEIS LTDA", "cnpj": "19.152.134/0001-56",
            "testemunha1_nome": "Jaime Perinazzo", "testemunha1_cpf": "123.456.789-00",
            "testemunha2_nome": "Felipe Guizalberte", "testemunha2_cpf": "987.654.321-00"}
    ctx = construir_contexto(
        cliente={"nome": "Ana Cliente", "cpf": "111.222.333-44", "email": "a@x.com",
                 "telefone": "(12) 9", "logradouro": "Rua A", "numero": "1", "complemento": "",
                 "bairro": "Centro", "cidade": "SJC", "cep": "12000", "estado": "SP",
                 "inst_mesmo_residencial": True, "inst_logradouro": "", "inst_numero": "",
                 "inst_complemento": "", "inst_bairro": "", "inst_cidade": "", "inst_cep": "", "inst_uf": ""},
        usuario={"nome": "Consultor Z", "telefone": "", "email": ""},
        forma_pagamento_json="",
        loja=loja,
    )
    path = preencher_contrato(91001, ctx)
    full = "\n".join(p.text for p in Document(path).paragraphs)
    os.remove(path)
    assert "Ana Cliente\n111.222.333-44" in full   # nome numa linha, CPF (valor) na linha de baixo
    assert "Consultor Z" in full             # consultor PERMANECE no cabeçalho (par. 0)
    assert "Jaime Perinazzo" in full
    assert "Felipe Guizalberte" in full


# NOTA: testes removidos nesta tarefa (comportamento intencionalmente eliminado,
# agora responsabilidade do modelo modelo_contrato_mapeado.docx):
#   - test_contrato_cpf_vira_cpf_cnpj: o relabel "CPF"->"CPF/CNPJ" era feito em
#     codigo (_relabel_cpf_cnpj, removido); o modelo ja traz "CPF/CNPJ" no texto.
#   - test_contrato_tags_nomenclatura: os rotulos cinza Pt-7 (_set_cell rotulo=)
#     foram removidos; o modelo ja contem os rotulos fixos das celulas.


def test_geracao_completa_sem_marcadores_remanescentes():
    import os, json, re
    from docx import Document
    from docx.oxml.ns import qn
    from mod_contrato import preencher_contrato, construir_contexto
    loja = {"nome": "INSPIRIUM MOVEIS LTDA", "cnpj": "19.152.134/0001-56",
            "testemunha1_nome": "Jaime Perinazzo", "testemunha1_cpf": "123.456.789-00",
            "testemunha2_nome": "Felipe Guizalberte", "testemunha2_cpf": "987.654.321-00"}
    ctx = construir_contexto(
        cliente={"nome": "Ana Cliente", "cpf": "111.222.333-44", "email": "a@x.com",
                 "telefone": "(12) 90000-0000", "logradouro": "Rua A", "numero": "10",
                 "complemento": "ap 1", "bairro": "Centro", "cidade": "SJC", "cep": "12000-000",
                 "estado": "SP", "inst_mesmo_residencial": True, "inst_logradouro": "",
                 "inst_numero": "", "inst_complemento": "", "inst_bairro": "", "inst_cidade": "",
                 "inst_cep": "", "inst_uf": ""},
        usuario={"nome": "Consultor Z", "telefone": "(12) 91111-1111", "email": "z@x.com"},
        forma_pagamento_json=json.dumps({
            "tipo": "aymore", "nome_forma": "Financiamento Aymoré",
            "entrada_valor": 20000, "entrada_data": "2026-06-18", "entrada_forma": "pix",
            "total_cliente": 129572.01, "texto_cartao": "",
            "parcelas": [{"num": i+1, "data": f"18/{7+i:02d}/2026", "valor": 4820.0} for i in range(3)]}),
        loja=loja)
    ctx["num_contrato"]  = "INS-2026-06-17-009"
    ctx["data_contrato"] = "17/06/2026"
    path = preencher_contrato(92001, ctx)
    doc = Document(path)
    blob = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                blob += "\n" + c.text
    for sec in doc.sections:
        for h in (sec.header,):
            blob += "\n" + "\n".join(tt.text or "" for tt in h._element.iter(qn('w:t')))
    os.remove(path)
    sobra = re.findall(r'\[[A-Za-z0-9_ ]+\]', blob)
    assert sobra == [], f"marcadores não substituídos: {sobra}"
    assert "Ana Cliente" in blob
    assert "INS-2026-06-17-009" in blob and "17/06/2026" in blob
    assert "R$ 129.572,01" in blob
    assert "R$ 4.820,00" in blob
    assert "Jaime Perinazzo" in blob and "Felipe Guizalberte" in blob


# ── Número do contrato ─────────────────────────────────────────────────────────

def test_gerar_num_contrato_formato():
    from datetime import datetime
    from mod_contrato import gerar_num_contrato
    n = gerar_num_contrato([], "INS", data=datetime(2026, 6, 17))
    assert n == "INS-2026-06-17-001"


def test_gerar_num_contrato_sequencia_continua():
    from datetime import datetime
    from mod_contrato import gerar_num_contrato
    existentes = ["INS-2026-06-15-001", "INS-2026-06-16-002", "ORZ-2026-06-16-009"]
    n = gerar_num_contrato(existentes, "INS", data=datetime(2026, 6, 17))
    assert n == "INS-2026-06-17-003"


def test_gerar_num_contrato_loja_customizada():
    from datetime import datetime
    from mod_contrato import gerar_num_contrato
    n = gerar_num_contrato([], "ORZ", data=datetime(2026, 1, 5))
    assert n == "ORZ-2026-01-05-001"


# ── Valores das parcelas ───────────────────────────────────────────────────────

def test_parse_pagamento_valores_alinhados():
    from mod_contrato import _parse_pagamento
    pag = json.dumps({
        "tipo": "aymore", "nome_forma": "Aymoré",
        "parcelas": [
            {"seq": 1, "data": "2026-07-10", "valor": "R$ 100,00"},
            {"seq": 2, "data": "2026-08-10", "valor": "R$ 200,00"},
        ],
    })
    d = _parse_pagamento(pag)
    assert d["num_parcelas_int"] == 2
    assert d["valores"][0] == "R$ 100,00"
    assert d["valores"][1] == "R$ 200,00"
    assert d["valores"][2] == ""        # preenchido até 24
    assert len(d["valores"]) == 24


def test_parse_pagamento_cartao_sem_valores():
    from mod_contrato import _parse_pagamento
    d = _parse_pagamento(json.dumps({"tipo": "cartao", "nome_forma": "Cartão"}))
    assert d["valores"] == [""] * 24
    assert d["num_parcelas_int"] == 0


# ── Grade de parcelas: ordinais, valores, traços, linhas removidas ─────────────

def _ctx_parcelas(n):
    from mod_contrato import construir_contexto
    parcelas = [{"seq": i + 1, "data": f"2026-{7+i:02d}-10", "valor": f"R$ {i+1}00,00"}
                for i in range(n)]
    return construir_contexto(
        cliente={"nome": "Ana", "cpf": "1", "email": "a@x.com", "telefone": "(12)9",
                 "logradouro": "Rua A", "numero": "10", "complemento": "", "bairro": "Centro",
                 "cidade": "SJC", "cep": "12000", "estado": "SP", "inst_mesmo_residencial": True,
                 "inst_logradouro": "", "inst_numero": "", "inst_complemento": "",
                 "inst_bairro": "", "inst_cidade": "", "inst_cep": "", "inst_uf": ""},
        usuario={"nome": "Y", "telefone": "", "email": ""},
        forma_pagamento_json=json.dumps({"tipo": "aymore", "nome_forma": "Aymoré",
                                          "parcelas": parcelas}))


# NOTA: testes removidos nesta tarefa (comportamento intencionalmente eliminado):
#   - test_grade_ordinais_valores_tracos_e_linhas_removidas: a grade nao usa mais
#     ordinais ("Nª") e nao remove linhas; _preencher_grade preenche por posicao
#     (valor+data, _TRACO nos slots vazios) preservando as 11 linhas da tabela.
#   - test_grade_a_vista_remove_todas_as_linhas: idem; linhas nao sao mais removidas.
#     Cobertura atual: test_preencher_grade_* e test_geracao_completa_*.


# ── Cabeçalho: num_contrato e data_contrato ───────────────────────────────────

def test_cabecalho_num_contrato_substituido():
    import os
    from mod_contrato import preencher_contrato, _MODELO
    if not os.path.exists(_MODELO):
        return
    from docx import Document
    from docx.oxml.ns import qn
    ctx = _ctx_parcelas(2)
    ctx["num_contrato"]  = "INS-2026-06-17-007"
    ctx["data_contrato"] = "17/06/2026"
    path = preencher_contrato(91006, ctx)
    doc = Document(path)
    hdr_text = []
    for sec in doc.sections:
        for h in (sec.header, sec.first_page_header, sec.even_page_header):
            hdr_text += [t.text for t in h._element.iter(qn('w:t')) if t.text]
    os.remove(path)
    blob = " ".join(hdr_text)
    assert "INS-2026-06-17-007" in blob
    assert "17/06/2026" in blob
    assert "[Num_Contrato]" not in blob
    assert "Data_contrato" not in blob


def test_template_oficial_tem_marcadores():
    import os
    from docx import Document
    from mod_contrato import _MODELO
    assert os.path.basename(_MODELO) == "modelo_contrato_mapeado.docx"
    assert os.path.exists(_MODELO)
    d = Document(_MODELO)
    blob = "\n".join(p.text for p in d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                blob += "\n" + c.text
    assert "[NOME_CLIENTE]" in blob
    assert "[TOTAL_CONTRATO]" in blob
    assert "[DATA_PARCELA_1]" in blob
    assert "[VALOR_PARCELA]" in blob


def test_parse_pagamento_estrutura_real():
    import json
    from mod_contrato import _parse_pagamento
    pag = json.dumps({
        "tipo": "aymore", "nome_forma": "Financiamento Aymoré",
        "entrada_valor": 20000, "entrada_data": "2026-06-18", "entrada_forma": "pix",
        "total_cliente": 129572.01, "texto_cartao": "",
        "parcelas": [
            {"num": 1, "data": "18/07/2026", "valor": 4820.00},
            {"num": 2, "data": "17/08/2026", "valor": 4820.00},
        ],
    })
    d = _parse_pagamento(pag)
    assert d["num_parcelas_int"] == 2
    assert d["valores"][0] == "R$ 4.820,00"
    assert d["valores"][1] == "R$ 4.820,00"
    assert d["valores"][2] == ""
    assert d["datas"][0] == "18/07/2026"
    assert d["datas"][2] == ""
    assert d["valor_contrato"] == "R$ 129.572,01"
    assert len(d["valores"]) == 24 and len(d["datas"]) == 24


def test_parse_pagamento_cartao_texto():
    import json
    from mod_contrato import _parse_pagamento
    d = _parse_pagamento(json.dumps({
        "tipo": "cartao", "nome_forma": "Cartão de Crédito",
        "texto_cartao": "12x R$ 10.000,00", "total_cliente": 120000, "parcelas": []}))
    assert d["texto_cartao"] == "12x R$ 10.000,00"
    assert d["num_parcelas_int"] == 0
    assert d["valores"] == [""] * 24
    assert d["valor_contrato"] == "R$ 120.000,00"


def test_substituir_marcadores_basico():
    from docx import Document
    from mod_contrato import _substituir_marcadores
    d = Document()
    d.add_paragraph("Cliente: [NOME_CLIENTE] CPF/CNPJ: [CPF]")
    d.add_paragraph("Desconhecido: [NAO_EXISTE]")
    _substituir_marcadores(d, {"NOME_CLIENTE": "Ana Lima", "CPF": "111.222.333-44"})
    txt = "\n".join(p.text for p in d.paragraphs)
    assert "Cliente: Ana Lima CPF/CNPJ: 111.222.333-44" in txt
    assert "[NAO_EXISTE]" in txt          # desconhecido permanece


def test_substituir_marcadores_case_e_duplo_colchete():
    from docx import Document
    from mod_contrato import _substituir_marcadores
    d = Document()
    d.add_paragraph("N: [Num_Contrato]  D: [[Data_contrato]")
    _substituir_marcadores(d, {"NUM_CONTRATO": "INS-2026-06-17-001", "DATA_CONTRATO": "17/06/2026"})
    txt = d.paragraphs[0].text
    assert "INS-2026-06-17-001" in txt and "17/06/2026" in txt
    assert "[" not in txt


def test_substituir_marcadores_em_tabela():
    from docx import Document
    from mod_contrato import _substituir_marcadores
    d = Document()
    t = d.add_table(rows=1, cols=1)
    t.rows[0].cells[0].paragraphs[0].add_run("Nome\n[NOME_CLIENTE]")
    _substituir_marcadores(d, {"NOME_CLIENTE": "Bia"})
    assert "Bia" in t.rows[0].cells[0].text
    assert "[NOME_CLIENTE]" not in t.rows[0].cells[0].text


def test_subst_preserva_formatacao_por_run():
    """Substituição preserva a formatação de CADA run: rótulo cinza pequeno
    continua cinza pequeno; valor em negrito 8.5 continua negrito 8.5."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from mod_contrato import _substituir_marcadores
    d = Document()
    par = d.add_paragraph()
    r1 = par.add_run("Nome\n")
    r1.font.size = Pt(7); r1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r2 = par.add_run("[NOME_CLIENTE]")
    r2.bold = True; r2.font.size = Pt(8.5)
    _substituir_marcadores(d, {"NOME_CLIENTE": "Ana Paula"})
    runs = d.paragraphs[0].runs
    rot = next(r for r in runs if "Nome" in r.text)
    val = next(r for r in runs if "Ana Paula" in r.text)
    # rótulo mantém cinza e tamanho pequeno
    assert str(rot.font.color.rgb) == "888888"
    assert rot.font.size.pt == 7.0
    # valor mantém negrito e tamanho
    assert val.bold is True
    assert val.font.size.pt == 8.5


# ── Grade de parcelas por posição (valor+data, traços, cartão) ─────────────────

def test_preencher_grade_valores_datas_e_tracos():
    from docx import Document
    from mod_contrato import _MODELO, _preencher_grade, _TRACO, _localizar_tabela
    d = Document(_MODELO)
    pag = {"tipo": "aymore", "num_parcelas_int": 2,
           "valores": ["R$ 4.820,00", "R$ 4.820,00"] + [""] * 22,
           "datas":   ["18/07/2026", "17/08/2026"] + [""] * 22,
           "texto_cartao": ""}
    _preencher_grade(d, pag)
    t3 = _localizar_tabela(d, "forma de pagamento")
    blob = " ".join(c.text for row in t3.rows for c in row.cells)
    assert "R$ 4.820,00" in blob
    assert "18/07/2026" in blob and "17/08/2026" in blob
    assert _TRACO in blob                       # slot 3 (vazio) na linha usada
    assert "[VALOR_PARCELA]" not in blob
    assert "[DATA_PARCELA_3]" not in blob
    # 2 parcelas → 1 linha de grade; linhas vazias eliminadas: 3 cabeçalho + 1 = 4
    assert len(t3.rows) == 4


def test_preencher_grade_cartao_valores_sem_data():
    from docx import Document
    from mod_contrato import _MODELO, _preencher_grade, _TRACO, _unique_cells, _localizar_tabela
    d = Document(_MODELO)
    _preencher_grade(d, {"tipo": "cartao", "num_parcelas_int": 4,
                         "valores": ["R$ 100,00"] * 4 + [""] * 20,
                         "datas": [""] * 24, "texto_cartao": "12x R$ 300,00"})
    t3 = _localizar_tabela(d, "forma de pagamento")
    # 4 parcelas → 2 linhas de grade; linhas vazias eliminadas: 3 cabeçalho + 2 = 5
    assert len(t3.rows) == 5
    # 1ª linha da grade: parcelas 1-3 → valor preenchido, data EM BRANCO
    cells = _unique_cells(t3.rows[3])
    for vcol, dcol in [(0, 1), (2, 3), (4, 5)]:
        assert cells[vcol].text == "R$ 100,00"
        assert cells[dcol].text == ""        # cartão: parcela sem data
    # 2ª linha: parcela 4 preenchida; slots 5,6 vazios → traços (valor e data)
    cells2 = _unique_cells(t3.rows[4])
    assert cells2[0].text == "R$ 100,00" and cells2[1].text == ""
    assert cells2[2].text == _TRACO and cells2[3].text == _TRACO
    assert cells2[4].text == _TRACO and cells2[5].text == _TRACO
    # texto_cartao NÃO é despejado na grade
    blob = " ".join(c.text for row in t3.rows for c in row.cells)
    assert "12x R$ 300,00" not in blob


def test_parse_pagamento_cartao_avista_num_parcelas():
    import json
    from mod_contrato import _parse_pagamento
    d = _parse_pagamento(json.dumps({"tipo": "cartao", "nome_forma": "Cartão de Crédito",
        "total_cliente": 5000, "parcelas": [{"num": 1, "valor": 5000, "data": ""}]}))
    assert d["num_parcelas"] == "à vista"
    assert d["num_parcelas_int"] == 1


def test_parse_pagamento_cartao_parcelado_num_parcelas_e_datas_vazias():
    import json
    from mod_contrato import _parse_pagamento
    d = _parse_pagamento(json.dumps({"tipo": "cartao", "nome_forma": "Cartão de Crédito",
        "total_cliente": 12000, "parcelas": [{"num": i+1, "valor": 1000, "data": ""} for i in range(12)]}))
    assert d["num_parcelas"] == "12"
    assert d["num_parcelas_int"] == 12
    assert all(x == "" for x in d["datas"][:12])   # cartão: parcelas sem data


def test_protegido_tem_documentprotection_e_regioes():
    import json, os
    from docx import Document
    from docx.oxml.ns import qn
    from mod_contrato import preencher_contrato, construir_contexto
    loja = {"nome": "LOJA TESTE", "cnpj": "00.000.000/0001-00",
            "testemunha1_nome": "T1", "testemunha1_cpf": "111.111.111-11",
            "testemunha2_nome": "T2", "testemunha2_cpf": "222.222.222-22"}
    ctx = construir_contexto(
        cliente={"nome":"Ana","cpf":"111","email":"a@x.com","telefone":"(12)9","logradouro":"R",
                 "numero":"1","complemento":"","bairro":"C","cidade":"SJC","cep":"1","estado":"SP",
                 "inst_mesmo_residencial":True,"inst_logradouro":"","inst_numero":"","inst_complemento":"",
                 "inst_bairro":"","inst_cidade":"","inst_cep":"","inst_uf":""},
        usuario={"nome":"Z","telefone":"","email":""},
        forma_pagamento_json=json.dumps({"tipo":"aymore","nome_forma":"Aymoré","total_cliente":1000,
            "texto_cartao":"","parcelas":[{"num":1,"data":"18/07/2026","valor":500.0}]}),
        loja=loja)
    ctx["num_contrato"]="INS-1"; ctx["data_contrato"]="18/06/2026"
    p = preencher_contrato(97001, ctx, protegido=True)
    d = Document(p)
    prot = d.settings.element.find(qn('w:documentProtection'))
    body_xml = d.element.body.xml
    os.remove(p)
    assert prot is not None and prot.get(qn('w:edit')) == "readOnly"
    assert "permStart" in body_xml and "permEnd" in body_xml
    assert "Ana" in body_xml

def test_nao_protegido_sem_documentprotection():
    import json, os
    from docx import Document
    from docx.oxml.ns import qn
    from mod_contrato import preencher_contrato, construir_contexto
    loja = {"nome": "LOJA TESTE", "cnpj": "00.000.000/0001-00",
            "testemunha1_nome": "T1", "testemunha1_cpf": "111.111.111-11",
            "testemunha2_nome": "T2", "testemunha2_cpf": "222.222.222-22"}
    ctx = construir_contexto(
        cliente={"nome":"Ana","cpf":"1","email":"","telefone":"","logradouro":"","numero":"",
                 "complemento":"","bairro":"","cidade":"","cep":"","estado":"","inst_mesmo_residencial":True,
                 "inst_logradouro":"","inst_numero":"","inst_complemento":"","inst_bairro":"",
                 "inst_cidade":"","inst_cep":"","inst_uf":""},
        usuario={"nome":"Z","telefone":"","email":""},
        forma_pagamento_json=json.dumps({"tipo":"aymore","nome_forma":"Aymoré","total_cliente":0,
            "texto_cartao":"","parcelas":[]}),
        loja=loja)
    p = preencher_contrato(97002, ctx, protegido=False)
    d = Document(p)
    has = d.settings.element.find(qn('w:documentProtection')) is not None
    body = d.element.body.xml
    os.remove(p)
    assert has is False
    assert "permStart" not in body

def test_protegido_mantem_texto_e_valores():
    # proteção não altera o conteúdo: mesmos valores que protegido=False
    import json, os, re
    from docx import Document
    from docx.oxml.ns import qn
    from mod_contrato import preencher_contrato, construir_contexto
    _loja = {"nome": "LOJA TESTE", "cnpj": "00.000.000/0001-00",
             "testemunha1_nome": "T1", "testemunha1_cpf": "111.111.111-11",
             "testemunha2_nome": "T2", "testemunha2_cpf": "222.222.222-22"}
    def gen(protegido):
        ctx = construir_contexto(
            cliente={"nome":"Ana Cliente","cpf":"111.222.333-44","email":"a@x.com","telefone":"(12)9",
                     "logradouro":"Rua A","numero":"10","complemento":"","bairro":"Centro","cidade":"SJC",
                     "cep":"12000","estado":"SP","inst_mesmo_residencial":True,"inst_logradouro":"",
                     "inst_numero":"","inst_complemento":"","inst_bairro":"","inst_cidade":"","inst_cep":"","inst_uf":""},
            usuario={"nome":"Z","telefone":"(12)9","email":"z@x.com"},
            forma_pagamento_json=json.dumps({"tipo":"aymore","nome_forma":"Aymoré","total_cliente":129572.01,
                "texto_cartao":"","parcelas":[{"num":i+1,"data":f"18/{7+i:02d}/2026","valor":4820.0} for i in range(3)]}),
            loja=_loja)
        ctx["num_contrato"]="INS-9"; ctx["data_contrato"]="18/06/2026"
        p = preencher_contrato(97003, ctx, protegido=protegido)
        d = Document(p)
        blob = "\n".join(par.text for par in d.paragraphs)
        for t in d.tables:
            for row in t.rows:
                for c in row.cells: blob += "\n"+c.text
        os.remove(p)
        return blob
    a = gen(True); b = gen(False)
    assert "Ana Cliente" in a and "R$ 4.820,00" in a and "R$ 129.572,01" in a
    assert sorted(re.findall(r'\[[A-Za-z0-9_ ]+\]', a)) == []   # nenhum marcador sobra
    assert a == b   # proteção não muda o texto


def test_converter_pdf_nao_regenera_docx(monkeypatch):
    import mod_contrato
    chamou = {"preencher": False, "convert_path": None}
    monkeypatch.setattr(mod_contrato, "preencher_contrato",
                        lambda *a, **k: chamou.__setitem__("preencher", True) or "X")
    def fake_run(cmd, **kw):
        chamou["convert_path"] = cmd[-1]
        class R: pass
        return R()
    monkeypatch.setattr(mod_contrato.subprocess, "run", fake_run)
    out = mod_contrato._converter_pdf("/tmp/contrato_5.docx")
    assert chamou["preencher"] is False
    assert chamou["convert_path"] == "/tmp/contrato_5.docx"
    assert out.endswith("contrato_5.pdf")


# ── Forma de pagamento: rótulos pt-BR + forma_parcela + marcador TIPO ──────────

def test_forma_label_mapeia_codigos():
    from mod_contrato import _forma_label
    assert _forma_label("pix") == "Pix"
    assert _forma_label("ted") == "TED"
    assert _forma_label("transferencia") == "TED"
    assert _forma_label("boleto") == "Boleto"
    assert _forma_label("cheque") == "Cheque"
    assert _forma_label("dinheiro") == "Dinheiro"
    assert _forma_label("cartao_credito") == "Cartão de Crédito"
    assert _forma_label("") == ""
    assert _forma_label("Boleto") == "Boleto"   # já-rótulo passa adiante


def test_parse_pagamento_forma_parcela_de_parcelas():
    import json
    from mod_contrato import _parse_pagamento
    pag = json.dumps({
        "tipo": "venda_programada", "nome_forma": "Venda Programada",
        "entrada_valor": 1000, "entrada_forma": "pix", "total_cliente": 5000,
        "parcelas": [
            {"num": 1, "data": "10/07/2026", "valor": 2000.0, "forma": "cheque"},
            {"num": 2, "data": "10/08/2026", "valor": 2000.0, "forma": "cheque"},
        ],
    })
    d = _parse_pagamento(pag)
    assert d["entrada_tipo"] == "Pix"
    assert d["forma_parcela"] == "Cheque"


def test_parse_pagamento_forma_parcela_cartao():
    import json
    from mod_contrato import _parse_pagamento
    d = _parse_pagamento(json.dumps({
        "tipo": "cartao", "nome_forma": "Cartão de Crédito",
        "texto_cartao": "12x R$ 10.000,00", "total_cliente": 120000, "parcelas": []}))
    assert d["forma_parcela"] == "Cartão de Crédito"


def test_montar_mapping_inclui_tipo():
    from mod_contrato import _montar_mapping
    ctx = {}
    pag = {"forma_parcela": "Boleto", "entrada_tipo": "Pix", "num_parcelas": "3"}
    m = _montar_mapping(ctx, pag)
    assert m["TIPO"] == "Boleto"
    assert m["FORMA_ENTRADA"] == "Pix"


def test_template_tem_marcador_tipo():
    import os
    from docx import Document
    from mod_contrato import _MODELO
    assert os.path.exists(_MODELO)
    d = Document(_MODELO)
    blob = "\n".join(p.text for p in d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                blob += "\n" + c.text
    assert "[TIPO]" in blob
    assert "[NUM_PARCELAS]" in blob


def test_geracao_completa_com_forma_parcela():
    import os, json, re
    from docx import Document
    from mod_contrato import preencher_contrato, construir_contexto
    loja = {"nome": "LOJA TESTE", "cnpj": "00.000.000/0001-00",
            "testemunha1_nome": "T1", "testemunha1_cpf": "111.111.111-11",
            "testemunha2_nome": "T2", "testemunha2_cpf": "222.222.222-22"}
    ctx = construir_contexto(
        cliente={"nome": "Ana", "cpf": "1", "email": "a@x.com", "telefone": "(12)9",
                 "logradouro": "Rua A", "numero": "10", "complemento": "", "bairro": "Centro",
                 "cidade": "SJC", "cep": "12000", "estado": "SP", "inst_mesmo_residencial": True,
                 "inst_logradouro": "", "inst_numero": "", "inst_complemento": "",
                 "inst_bairro": "", "inst_cidade": "", "inst_cep": "", "inst_uf": ""},
        usuario={"nome": "Z", "telefone": "", "email": ""},
        forma_pagamento_json=json.dumps({
            "tipo": "venda_programada", "nome_forma": "Venda Programada",
            "entrada_valor": 1000, "entrada_data": "2026-06-18", "entrada_forma": "pix",
            "total_cliente": 5000.0, "texto_cartao": "",
            "parcelas": [{"num": i+1, "data": f"18/{7+i:02d}/2026", "valor": 2000.0,
                          "forma": "cheque"} for i in range(2)]}),
        loja=loja)
    ctx["num_contrato"] = "INS-2026-06-18-001"; ctx["data_contrato"] = "18/06/2026"
    path = preencher_contrato(93001, ctx)
    doc = Document(path)
    blob = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                blob += "\n" + c.text
    os.remove(path)
    assert "Cheque" in blob
    assert "Pix" in blob
    assert re.findall(r'\[[A-Za-z0-9_ ]+\]', blob) == []


# ── Assinaturas: formatação uniforme do nome do cliente ───────────────────────

def test_template_assinaturas_normalizadas():
    from docx import Document
    from mod_contrato import _MODELO
    d = Document(_MODELO)
    alvos = {"[NOME_CLIENTE]": False, "[NOME_TESTEMUNHA_2]": False}
    for p in d.paragraphs:
        for marc in list(alvos):
            if marc in p.text:
                alvos[marc] = True
                assert p.style.name == "Heading 2", f"{marc} estilo={p.style.name}"
                # sem run inicial vazio/quebra de linha
                assert p.runs and (p.runs[0].text or "").strip() != ""
    assert all(alvos.values()), f"marcadores não encontrados: {alvos}"


def test_assinatura_cliente_mesmo_estilo_da_empresa():
    import os
    from docx import Document
    from mod_contrato import preencher_contrato, construir_contexto
    loja = {"nome": "INSPIRIUM MOVEIS PLANEJADOS E DECORACAO LTDA", "cnpj": "19.152.134/0001-56",
            "testemunha1_nome": "Jaime", "testemunha1_cpf": "123.456.789-00",
            "testemunha2_nome": "Felipe", "testemunha2_cpf": "987.654.321-00"}
    ctx = construir_contexto(
        cliente={"nome": "Ana Cliente", "cpf": "111.222.333-44", "email": "a@x.com",
                 "telefone": "(12)9", "logradouro": "Rua A", "numero": "10",
                 "complemento": "", "bairro": "Centro", "cidade": "SJC", "cep": "12000",
                 "estado": "SP", "inst_mesmo_residencial": True, "inst_logradouro": "",
                 "inst_numero": "", "inst_complemento": "", "inst_bairro": "",
                 "inst_cidade": "", "inst_cep": "", "inst_uf": ""},
        usuario={"nome": "Z", "telefone": "", "email": ""},
        forma_pagamento_json="",
        loja=loja)
    path = preencher_contrato(94001, ctx)
    d = Document(path)
    estilo_cliente = estilo_empresa = None
    for p in d.paragraphs:
        if "Ana Cliente" in p.text:
            estilo_cliente = p.style.name
        if "INSPIRIUM MOVEIS PLANEJADOS" in p.text:
            estilo_empresa = p.style.name
    os.remove(path)
    assert estilo_cliente == "Heading 2"
    assert estilo_empresa == "Heading 2"
    assert estilo_cliente == estilo_empresa


def test_assinaturas_nome_e_cpf_em_linhas_separadas():
    """Cada signatário do bloco de assinatura tem o NOME numa linha e o
    marcador de CPF/CNPJ na linha imediatamente abaixo (nova estrutura)."""
    from docx import Document
    from mod_contrato import _MODELO
    d = Document(_MODELO)
    pars = [(p.text or "").strip() for p in d.paragraphs]

    def linha_seguinte(marcador_nome, marcador_doc):
        for i, t in enumerate(pars):
            if marcador_nome in t:
                assert "CPF" not in t and "CNPJ" not in t, f"{marcador_nome} tem doc na mesma linha: {t!r}"
                j = i + 1
                while j < len(pars) and not pars[j]:
                    j += 1
                assert j < len(pars) and marcador_doc in pars[j], \
                    f"linha de doc de {marcador_nome} inesperada: {pars[j] if j < len(pars) else None!r}"
                return
        raise AssertionError(f"marcador {marcador_nome} não encontrado")

    linha_seguinte("[NOME_EMPRESA]",      "[CNPJ_EMPRESA]")
    linha_seguinte("[NOME_CLIENTE]",      "[CPF_CLIENTE]")
    linha_seguinte("[NOME_TESTEMUNHA_1]", "[CPF_TESTEMUNHA_1]")
    linha_seguinte("[NOME_TESTEMUNHA_2]", "[CPF_TESTEMUNHA_2]")


def test_montar_mapping_inclui_empresa_e_cpfs():
    from mod_contrato import _montar_mapping
    loja = {"nome": "INSPIRIUM MOVEIS LTDA", "cnpj": "19.152.134/0001-56",
            "testemunha1_nome": "Jaime", "testemunha1_cpf": "123.456.789-00",
            "testemunha2_nome": "Felipe", "testemunha2_cpf": "987.654.321-00"}
    ctx = {"cliente_cpf": "111.222.333-44", "loja": loja}
    m = _montar_mapping(ctx, {})
    assert m["NOME_EMPRESA"] == "INSPIRIUM MOVEIS LTDA"
    assert m["CNPJ_EMPRESA"] == "19.152.134/0001-56"
    assert m["CPF_CLIENTE"] == "111.222.333-44"
    assert m["CPF_TESTEMUNHA_1"] == "123.456.789-00"
    assert m["CPF_TESTEMUNHA_2"] == "987.654.321-00"
    assert m["NOME_TESTEMUNHA_1"] == "Jaime"
    assert m["NOME_TESTEMUNHA_2"] == "Felipe"


def test_substituir_marcadores_cabecalho_fragmentado():
    """Marcador fragmentado em múltiplos runs no cabeçalho é substituído."""
    from docx import Document
    from docx.oxml.ns import qn
    from mod_contrato import _substituir_marcadores
    d = Document()
    hdr = d.sections[0].header
    p = hdr.paragraphs[0]
    p.add_run("[")
    p.add_run("NUM_CONTRATO")
    p.add_run("]")
    _substituir_marcadores(d, {"NUM_CONTRATO": "INS-2026-01-01-001"})
    txt = "".join(t.text or "" for t in hdr._element.iter(qn('w:t')))
    assert "INS-2026-01-01-001" in txt
    assert "[NUM_CONTRATO]" not in txt


# ── Valor por ambiente no contrato (rateio do financeiro) ──────────────────────

def test_ambientes_valor_proporcional_ao_vava():
    from mod_contrato import ambientes_valor_contrato
    out = ambientes_valor_contrato([("Cozinha", 100.0), ("Sala", 300.0)],
                                   vavo=400.0, val_cont=440.0)
    assert out == [("Cozinha", 110.0), ("Sala", 330.0)]
    assert round(sum(v for _, v in out), 2) == 440.0


def test_ambientes_reconciliacao_residuo_no_ultimo():
    from mod_contrato import ambientes_valor_contrato
    out = ambientes_valor_contrato(
        [("A", 33.33), ("B", 33.33), ("C", 33.34)], vavo=100.0, val_cont=100.01)
    assert round(sum(v for _, v in out), 2) == 100.01
    # o resíduo de arredondamento cai no último ambiente
    assert out[-1][0] == "C"
    assert out[0][1] == 33.33 and out[1][1] == 33.33


def test_ambientes_sem_financeiro_valor_igual_vava():
    from mod_contrato import ambientes_valor_contrato
    out = ambientes_valor_contrato([("A", 100.0), ("B", 300.0)],
                                   vavo=400.0, val_cont=400.0)
    assert out == [("A", 100.0), ("B", 300.0)]


def test_ambientes_vavo_zero_nao_divide():
    from mod_contrato import ambientes_valor_contrato
    out = ambientes_valor_contrato([("A", 0.0), ("B", 0.0)], vavo=0.0, val_cont=0.0)
    assert out == [("A", 0.0), ("B", 0.0)]


def test_ambientes_lista_vazia():
    from mod_contrato import ambientes_valor_contrato
    assert ambientes_valor_contrato([], vavo=0.0, val_cont=0.0) == []


def test_localizar_tabela_forma_pagamento():
    from docx import Document
    from mod_contrato import _localizar_tabela, _MODELO
    doc = Document(_MODELO)
    t = _localizar_tabela(doc, "forma de pagamento")
    assert t is not None
    cab = " ".join(c.text for c in t.rows[0].cells).lower()
    assert "forma de pagamento" in cab


def test_preencher_ambientes_clona_linhas_do_template():
    """Preenche a tabela 'Ambientes do Projeto' do template (2 por linha),
    clonando linhas completas para ambientes extras e sem criar tabela nova."""
    from docx import Document
    from mod_contrato import _preencher_ambientes, _localizar_tabela, _unique_cells, _TRACO, _MODELO
    doc = Document(_MODELO)
    itens = [("Cozinha", 12345.67), ("Dormitório", 8900.0), ("Home Theater", 5200.0)]
    _preencher_ambientes(doc, itens)

    t = _localizar_tabela(doc, "ambientes do projeto")
    assert t is not None
    # header + ceil(3/2)=2 linhas de dados + linha de total = 4 linhas
    assert len(t.rows) == 4
    blob = "\n".join(c.text for r in t.rows for c in r.cells)
    # nomes e valores presentes
    for nome in ("Cozinha", "Dormitório", "Home Theater"):
        assert nome in blob
    for val in ("R$ 12.345,67", "R$ 8.900,00", "R$ 5.200,00"):
        assert val in blob
    # nenhum marcador de ambiente sobrou
    assert "NOME_AMBIENTE" not in blob and "VALOR_AMBIENTE" not in blob
    # cada linha de dados é COMPLETA (4 células únicas)
    for row in (t.rows[1], t.rows[2]):
        assert len(_unique_cells(row)) == 4
    # rótulos "Ambiente"/"Valor" das células são PRESERVADOS (só o marcador é trocado)
    prim = _unique_cells(t.rows[1])
    assert "Ambiente" in prim[0].text and "Valor" in prim[1].text
    assert "Cozinha" in prim[0].text                 # rótulo + valor coexistem
    # linha ímpar: 2ª metade da última linha de dados recebe traços (como na grade)
    ult = _unique_cells(t.rows[2])
    assert "Home Theater" in ult[0].text
    assert _TRACO in ult[2].text and _TRACO in ult[3].text
    assert "Ambiente" in ult[2].text and "Valor" in ult[3].text  # rótulo mantido no traço
    # NÃO cria tabela duplicada de ambientes
    n_amb = sum(1 for tb in doc.tables
                if "ambiente" in " ".join(c.text for c in tb.rows[0].cells).lower())
    assert n_amb == 1


def test_preencher_ambientes_par_uma_linha():
    """Nº par de ambientes → uma linha de dados por par (sem sobra)."""
    from docx import Document
    from mod_contrato import _preencher_ambientes, _localizar_tabela, _unique_cells, _MODELO
    doc = Document(_MODELO)
    _preencher_ambientes(doc, [("Cozinha", 100.0), ("Sala", 200.0)])
    t = _localizar_tabela(doc, "ambientes do projeto")
    assert len(t.rows) == 3  # header + 1 linha de dados + total
    cels = _unique_cells(t.rows[1])
    assert "Cozinha" in cels[0].text and "Sala" in cels[2].text
    assert "R$ 100,00" in cels[1].text and "R$ 200,00" in cels[3].text


def test_preencher_ambientes_lista_vazia_traces():
    """Lista vazia → linha-modelo com traços (nenhum marcador cru vaza)."""
    from docx import Document
    from mod_contrato import _preencher_ambientes, _localizar_tabela, _unique_cells, _TRACO, _MODELO
    doc = Document(_MODELO)
    _preencher_ambientes(doc, [])
    t = _localizar_tabela(doc, "ambientes do projeto")
    blob = "\n".join(c.text for r in t.rows for c in r.cells)
    assert "NOME_AMBIENTE" not in blob and "VALOR_AMBIENTE" not in blob
    # header + 1 linha de dados (traços) + total
    assert len(t.rows) == 3
    # marcadores viram traços; rótulos preservados
    dados = _unique_cells(t.rows[1])
    assert all(_TRACO in c.text for c in dados)
    assert "Ambiente" in dados[0].text and "Valor" in dados[1].text


def test_contrato_com_secao_ambientes():
    import os, json
    from docx import Document
    from mod_contrato import preencher_contrato, construir_contexto
    loja = {"nome": "INSPIRIUM MOVEIS LTDA", "cnpj": "19.152.134/0001-56",
            "testemunha1_nome": "Jaime", "testemunha1_cpf": "123.456.789-00",
            "testemunha2_nome": "Felipe", "testemunha2_cpf": "987.654.321-00"}
    ctx = construir_contexto(
        cliente={"nome": "Ana", "cpf": "111.222.333-44", "email": "a@x.com",
                 "telefone": "(12) 90000-0000", "logradouro": "Rua A", "numero": "10",
                 "complemento": "", "bairro": "Centro", "cidade": "SJC", "cep": "12000-000",
                 "estado": "SP", "inst_mesmo_residencial": True, "inst_logradouro": "",
                 "inst_numero": "", "inst_complemento": "", "inst_bairro": "", "inst_cidade": "",
                 "inst_cep": "", "inst_uf": ""},
        usuario={"nome": "Consultor Z", "telefone": "(12) 91111-1111", "email": "z@x.com"},
        forma_pagamento_json=json.dumps({
            "tipo": "aymore", "nome_forma": "Financiamento Aymoré",
            "entrada_valor": 20000, "entrada_data": "2026-06-18", "entrada_forma": "pix",
            "total_cliente": 26445.67, "texto_cartao": "",
            "parcelas": [{"num": i+1, "data": f"18/{7+i:02d}/2026", "valor": 4820.0} for i in range(3)]}),
        loja=loja)
    ctx["num_contrato"]  = "INS-2026-07-01-001"
    ctx["data_contrato"] = "01/07/2026"
    ctx["_ambientes"] = [("Cozinha", 12345.67), ("Dormitório", 8900.0), ("Home Theater", 5200.0)]
    path = preencher_contrato(93001, ctx)
    doc = Document(path)
    # coleta texto de todas as tabelas
    tbl_blob = ""
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                tbl_blob += "\n" + c.text
    # nenhum marcador de ambiente sobrou no documento
    n_amb_tabelas = sum(1 for t in doc.tables
                        if "ambiente" in " ".join(c.text for c in t.rows[0].cells).lower())
    os.remove(path)
    # ambientes preenchidos na tabela do template (nomes e valores)
    assert "Cozinha" in tbl_blob and "Dormitório" in tbl_blob and "Home Theater" in tbl_blob
    assert "R$ 12.345,67" in tbl_blob and "R$ 8.900,00" in tbl_blob and "R$ 5.200,00" in tbl_blob
    assert "NOME_AMBIENTE" not in tbl_blob and "VALOR_AMBIENTE" not in tbl_blob
    # total (VALOR DO CONTRATO) = soma, via marcador [TOTAL_CONTRATO]
    assert "R$ 26.445,67" in tbl_blob
    # não duplica a tabela de ambientes (a do template é a única)
    assert n_amb_tabelas == 1
    # Forma de Pagamento e grade ainda preenchidas
    assert "5. Forma de Pagamento" in tbl_blob
    assert "R$ 4.820,00" in tbl_blob
