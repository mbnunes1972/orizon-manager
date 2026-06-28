import os
from docx import Document

MARCADORES = ["NOME_EMPRESA", "NOME_CLIENTE", "CPF_CLIENTE", "AMBIENTES_LISTA",
              "VALOR_BRUTO", "DESCONTO_PCT", "VALOR_TOTAL", "MODALIDADE", "VALIDADE"]


def test_modelo_proposta_existe_e_tem_marcadores():
    path = os.path.join(os.path.dirname(__file__), "..", "modelo_proposta.docx")
    assert os.path.exists(path), "modelo_proposta.docx não foi gerado"
    doc = Document(path)
    texto = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                texto += "\n" + cell.text
    for m in MARCADORES:
        assert "[%s]" % m in texto, "marcador faltando: %s" % m
