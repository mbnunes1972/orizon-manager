"""Gera config/contrato_template.docx com todas as variáveis do sistema."""
import os
from docx import Document
from docx.shared import Pt

os.makedirs("config", exist_ok=True)

doc = Document()
doc.add_heading("CONTRATO DE COMPRA E VENDA", 0)

doc.add_paragraph("Cliente: {{ cliente_nome }}")
doc.add_paragraph("CPF: {{ cliente_cpf }}")
doc.add_paragraph("Endereço do Cliente: {{ cliente_endereco }}")
doc.add_paragraph("Telefone: {{ cliente_telefone }}")
doc.add_paragraph()
doc.add_paragraph("Endereço de Instalação: {{ endereco_instalacao }}")
doc.add_paragraph()
doc.add_heading("Projeto", level=1)
doc.add_paragraph("Projeto: {{ projeto_nome }}")
doc.add_paragraph("Data do Projeto: {{ projeto_data }}")
doc.add_paragraph("Orçamento: {{ orcamento_nome }}")
doc.add_paragraph()
doc.add_heading("Condições Comerciais", level=1)
doc.add_paragraph("Valor Total: {{ valor_total }}")
doc.add_paragraph("Forma de Pagamento: {{ forma_pagamento }}")
doc.add_paragraph("Entrada: {{ entrada_valor }}")
doc.add_paragraph("Parcelas: {{ parcelas_descricao }}")
doc.add_paragraph()
doc.add_heading("Ambientes", level=1)
doc.add_paragraph("{{ ambientes_lista }}")
doc.add_paragraph()
doc.add_heading("Assinaturas", level=1)
doc.add_paragraph("Consultor: {{ consultor_nome }}")
doc.add_paragraph("Data: {{ data_contrato }}")
doc.add_paragraph()
doc.add_heading("Adendo", level=1)
doc.add_paragraph("{{ adendo }}")

doc.save("config/contrato_template.docx")
print("Template criado em config/contrato_template.docx")
