"""Insere o marcador [TIPO] (forma das parcelas) junto de [NUM_PARCELAS] no
modelo_contrato_mapeado.docx. Idempotente: não duplica se já existir [TIPO].

Uso: python scripts/inserir_marcador_tipo.py
"""
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document  # noqa: E402

MODELO = os.path.join(os.path.dirname(__file__), "..", "modelo_contrato_mapeado.docx")


def _inserir_em_paragrafo(par):
    """Se o parágrafo contém [NUM_PARCELAS] e não contém [TIPO], reescreve seus
    runs inserindo ' / [TIPO]' logo após [NUM_PARCELAS], preservando a fonte do
    primeiro run. Retorna True se alterou."""
    txt = "".join(r.text for r in par.runs)
    if "[NUM_PARCELAS]" not in txt or "[TIPO]" in txt:
        return False
    novo = txt.replace("[NUM_PARCELAS]", "[NUM_PARCELAS] / [TIPO]")
    base = par.runs[0] if par.runs else None
    name = base.font.name if base is not None else None
    size = base.font.size if base is not None else None
    bold = base.bold if base is not None else None
    for r in list(par.runs):
        r._r.getparent().remove(r._r)
    run = par.add_run(novo)
    if name is not None:
        run.font.name = name
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    return True


def main():
    doc = Document(MODELO)
    alterou = False
    for par in doc.paragraphs:
        alterou = _inserir_em_paragrafo(par) or alterou
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    alterou = _inserir_em_paragrafo(par) or alterou
    if alterou:
        doc.save(MODELO)
        print("[OK] [TIPO] inserido junto de [NUM_PARCELAS].")
    else:
        print("[OK] Nada a fazer (já contém [TIPO] ou não achou [NUM_PARCELAS]).")


if __name__ == "__main__":
    main()
