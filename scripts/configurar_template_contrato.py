"""
Configura "Modelo de Contrato.docx" adicionando variáveis Jinja2 do sistema.
Execute a partir da raiz do projeto: python scripts/configurar_template_contrato.py
Salva em config/contrato_template.docx (template ativo do sistema).
"""
import os
from docx import Document
from docx.shared import Pt

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ORIGEM  = os.path.join(BASE, "Modelo de Contrato.docx")
DESTINO = os.path.join(BASE, "config", "contrato_template.docx")

doc = Document(ORIGEM)

# ── 1. Parágrafos 4-8 (capa — dados do cliente) ──────────────────────────────
# Esses parágrafos estão vazios na capa; vamos preenchê-los com variáveis.

def _set_par(p, texto, negrito=False, tamanho=None):
    """Limpa o parágrafo e adiciona um único run com o texto."""
    # Preserva o estilo mas limpa os runs
    for run in list(p.runs):
        run.text = ""
    # Usa o primeiro run existente ou cria um novo
    if p.runs:
        r = p.runs[0]
    else:
        r = p.add_run()
    r.text = texto
    if negrito:
        r.bold = True
    if tamanho:
        r.font.size = Pt(tamanho)

_set_par(doc.paragraphs[4], "CONTRATANTE: {{ cliente_nome }}", negrito=True)
_set_par(doc.paragraphs[5], "CPF: {{ cliente_cpf }}     Telefone: {{ cliente_telefone }}")
_set_par(doc.paragraphs[6], "Endereço do Contratante: {{ cliente_endereco }}")
_set_par(doc.paragraphs[7], "Endereço de instalação: {{ endereco_instalacao }}")
_set_par(doc.paragraphs[8], "Projeto: {{ projeto_nome }}    Orçamento: {{ orcamento_nome }}")

# ── 2. Tabela de valores e condição de pagamento ──────────────────────────────
tabela = doc.tables[0]

def _set_celula(cell, texto):
    """Substitui o conteúdo de uma célula por um único parágrafo/run."""
    p = cell.paragraphs[0]
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = texto
    else:
        p.add_run(texto)
    # Limpar parágrafos extras
    for p_extra in list(cell.paragraphs[1:]):
        p_elem = p_extra._element
        p_elem.getparent().remove(p_elem)

# Linha 1: projeto e valor
_set_celula(tabela.rows[1].cells[1], "{{ ambientes_lista }}")
_set_celula(tabela.rows[1].cells[2], "{{ valor_total }}")

# Linha 2: total
_set_celula(tabela.rows[2].cells[2], "Total: {{ valor_total }}")

# Linha 3: valor líquido da loja
_set_celula(tabela.rows[3].cells[2], "Valor líquido: {{ valor_liquido }}")

# Linha 4: condição de pagamento
_set_celula(tabela.rows[4].cells[0],
            "Entrada: {{ entrada_valor }}\n{{ parcelas_descricao }}")

# Limpar linhas de parcelas fixas (5-10) — substituir pela condição geral
for ri in range(5, len(tabela.rows)):
    for cell in tabela.rows[ri].cells:
        _set_celula(cell, "")

# ── 3. Data na área de assinatura ────────────────────────────────────────────
for p in doc.paragraphs:
    if "de junho de 2026" in p.text or (
        "Jos" in p.text and "Campos" in p.text and "202" in p.text
    ):
        _set_par(p, "São José dos Campos - SP, {{ data_contrato }}.")
        break

# ── 4. Nome do contratante na linha de assinatura ────────────────────────────
nome_count = 0
for p in doc.paragraphs:
    t = p.text.strip()
    if t == "NOME:" or t.startswith("NOME:"):
        if nome_count == 0:
            _set_par(p, "NOME: {{ cliente_nome }}")
            nome_count += 1
        elif nome_count == 1:
            _set_par(p, "NOME: ___________________________________")
            nome_count += 1

doc_count = 0
for p in doc.paragraphs:
    t = p.text.strip()
    if t == "Documento:" or t.startswith("Documento:"):
        if doc_count == 0:
            _set_par(p, "CPF: {{ cliente_cpf }}")
            doc_count += 1
        elif doc_count == 1:
            _set_par(p, "CPF: ___________________________________")
            doc_count += 1

# ── Salvar ────────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(DESTINO), exist_ok=True)
doc.save(DESTINO)
print(f"Template configurado e salvo em: {DESTINO}")
print("\nVariáveis inseridas:")
variaveis = [
    "{{ cliente_nome }}", "{{ cliente_cpf }}", "{{ cliente_telefone }}",
    "{{ cliente_endereco }}", "{{ endereco_instalacao }}",
    "{{ projeto_nome }}", "{{ orcamento_nome }}",
    "{{ ambientes_lista }}", "{{ valor_total }}", "{{ valor_liquido }}",
    "{{ entrada_valor }}", "{{ parcelas_descricao }}", "{{ data_contrato }}",
]
for v in variaveis:
    print(f"  {v}")
