"""Organiza o bloco de assinaturas do modelo_contrato_mapeado.docx:
cada signatário/testemunha fica com o NOME numa linha e o CPF/CNPJ na linha de
baixo, todos no mesmo estilo (Heading 2) e sem overrides de fonte (herdam o
estilo). Idempotente.

Trata: empresa (INSPIRIUM ... CNPJ:), cliente ([NOME_CLIENTE] CPF/CNPJ: [CPF]) e
testemunha 2 ([NOME_TESTEMUNHA_2] CPF: [CPF]). A testemunha 1 já está separada.

Uso: python scripts/organizar_assinaturas.py
"""
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document             # noqa: E402
from docx.oxml import OxmlElement     # noqa: E402
from docx.oxml.ns import qn           # noqa: E402
from docx.text.paragraph import Paragraph  # noqa: E402

MODELO = os.path.join(os.path.dirname(__file__), "..", "modelo_contrato_mapeado.docx")

_LABELS = ("CPF/CNPJ:", "CNPJ:", "CPF:")
_TOKENS = ("INSPIRIUM", "[NOME_CLIENTE]", "[NOME_TESTEMUNHA")


def _split_idx(text):
    """Índice do rótulo de CPF/CNPJ quando há NOME antes dele; -1 caso contrário."""
    idxs = [text.find(l) for l in _LABELS]
    idxs = [i for i in idxs if i > 0]   # i>0 garante que há nome antes do rótulo
    return min(idxs) if idxs else -1


def _limpar_rpr(run):
    rpr = run._r.find(qn("w:rPr"))
    if rpr is None:
        return
    for tag in ("w:b", "w:bCs", "w:rFonts", "w:sz", "w:szCs"):
        el = rpr.find(qn(tag))
        if el is not None:
            rpr.remove(el)


def _set_paragrafo(par, texto, estilo):
    """Deixa o parágrafo com um único run = texto, estilo dado, sem overrides."""
    for r in list(par.runs):
        r._r.getparent().remove(r._r)
    par.style = estilo
    run = par.add_run(texto)
    _limpar_rpr(run)


def main():
    doc = Document(MODELO)
    h2 = doc.styles["Heading 2"]
    # snapshot dos parágrafos-alvo antes de inserir novos
    alvos = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if len(t) >= 160:
            continue
        idx = _split_idx(t)
        if idx < 0:
            continue
        nome = t[:idx].strip()
        if not any(tok in nome for tok in _TOKENS):
            continue
        alvos.append(p)

    if not alvos:
        print("[OK] Nada a fazer (assinaturas já organizadas).")
        return

    for p in alvos:
        t = (p.text or "").strip()
        idx = _split_idx(t)
        nome = t[:idx].strip()
        cpf  = t[idx:].strip()
        # parágrafo atual = só o nome
        _set_paragrafo(p, nome, h2)
        # novo parágrafo logo abaixo = CPF/CNPJ
        novo_p = OxmlElement("w:p")
        p._p.addnext(novo_p)
        np = Paragraph(novo_p, p._parent)
        _set_paragrafo(np, cpf, h2)

    doc.save(MODELO)
    print(f"[OK] {len(alvos)} linha(s) de assinatura reorganizada(s) (nome / CPF).")


if __name__ == "__main__":
    main()
