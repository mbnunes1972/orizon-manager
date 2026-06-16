"""
preparar_template_contrato.py
Lê modelo_contrato_final.docx, insere placeholders Jinja2 nas células da capa,
salva como config/contrato_template.docx.
Execute após atualizar o modelo.
"""
import os
from docx import Document

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC  = os.path.join(BASE_DIR, "modelo_contrato_final.docx")
DEST = os.path.join(BASE_DIR, "config", "contrato_template.docx")


def _set_cell(cell, text: str):
    """Substitui o texto da célula pelo placeholder, preservando formatação do primeiro run."""
    para = cell.paragraphs[0]
    # Captura fonte/tamanho/bold do primeiro run (se existir)
    font_name, font_size, bold = None, None, None
    if para.runs:
        r0 = para.runs[0]
        font_name = r0.font.name
        font_size = r0.font.size
        bold = r0.bold
    # Limpa todos os runs existentes
    for run in para.runs:
        run.text = ""
    # Cria run com o placeholder
    run = para.add_run(text)
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size
    if bold is not None:
        run.bold = bold


def _unique_cells(row):
    """Retorna apenas as células únicas (não duplicadas por merge) de uma linha."""
    seen = set()
    result = []
    for c in row.cells:
        if id(c) not in seen:
            seen.add(id(c))
            result.append(c)
    return result


def main():
    os.makedirs(os.path.join(BASE_DIR, "config"), exist_ok=True)

    doc = Document(SRC)

    # ── Parágrafo 0: linha do consultor ──────────────────────────────────────
    # Estrutura real: 12 runs separando "Consultor:", tabs, "Telefone:", tabs, "e-mail:"
    # Estratégia: limpar todos os runs e reescrever o conteúdo completo no primeiro run
    p0 = doc.paragraphs[0]
    for run in p0.runs:
        run.text = ""
    if p0.runs:
        p0.runs[0].text = (
            "Consultor: {{ consultor_nome }}\t\t\t\t\t"
            "Telefone: {{ consultor_tel }}\t\t\t\t"
            "e-mail: {{ consultor_email }}"
        )
    else:
        p0.add_run(
            "Consultor: {{ consultor_nome }}\t\t\t\t\t"
            "Telefone: {{ consultor_tel }}\t\t\t\t"
            "e-mail: {{ consultor_email }}"
        )

    tables = doc.tables

    # ── Tabela 0: Identificação do cliente ────────────────────────────────────
    # R0: header (merged), R1: Nome | CPF, R2: E-mail | Telefone
    _set_cell(tables[0].rows[1].cells[0], "{{ cliente_nome }}")
    _set_cell(tables[0].rows[1].cells[1], "{{ cliente_cpf }}")
    _set_cell(tables[0].rows[2].cells[0], "{{ cliente_email }}")
    _set_cell(tables[0].rows[2].cells[1], "{{ cliente_telefone }}")

    # ── Tabela 1: Endereço residencial ────────────────────────────────────────
    # R0: header (merged), R1: Logradouro (fully merged — 1 unique cell)
    # R2: Número | Complemento | Bairro (3 unique cells)
    # R3: Cidade | CEP | UF (3 unique cells)
    _set_cell(_unique_cells(tables[1].rows[1])[0], "{{ res_logradouro }}")
    row2_t1 = _unique_cells(tables[1].rows[2])
    _set_cell(row2_t1[0], "{{ res_numero }}")
    _set_cell(row2_t1[1], "{{ res_complemento }}")
    _set_cell(row2_t1[2], "{{ res_bairro }}")
    row3_t1 = _unique_cells(tables[1].rows[3])
    _set_cell(row3_t1[0], "{{ res_cidade }}")
    _set_cell(row3_t1[1], "{{ res_cep }}")
    _set_cell(row3_t1[2], "{{ res_uf }}")

    # ── Tabela 2: Endereço de instalação ──────────────────────────────────────
    # Mesma estrutura da Tabela 1
    _set_cell(_unique_cells(tables[2].rows[1])[0], "{{ inst_logradouro }}")
    row2_t2 = _unique_cells(tables[2].rows[2])
    _set_cell(row2_t2[0], "{{ inst_numero }}")
    _set_cell(row2_t2[1], "{{ inst_complemento }}")
    _set_cell(row2_t2[2], "{{ inst_bairro }}")
    row3_t2 = _unique_cells(tables[2].rows[3])
    _set_cell(row3_t2[0], "{{ inst_cidade }}")
    _set_cell(row3_t2[1], "{{ inst_cep }}")
    _set_cell(row3_t2[2], "{{ inst_uf }}")

    # ── Tabela 3: Forma de pagamento ──────────────────────────────────────────
    # R0: header (fully merged)
    # R1: 3 pares merged (0+1=Entrada, 2+3=Tipo, 4+5=Data) — 3 unique cells at 0,2,4
    # R2: 3 pares merged (0+1=Modalidade, 2+3=Parcelas, 4+5=DataPrimeira) — 3 unique cells
    # R3–R10: 6 células independentes (1x, Data, 2x, Data, 3x, Data) por linha
    t3 = doc.tables[3]
    t3_rows = t3.rows

    # Linha de entrada (R1): unique cells são 0, 2, 4
    row1_unique = _unique_cells(t3_rows[1])
    # row1_unique[0]=Entrada(label+valor), [1]=Tipo(label+valor), [2]=Data(label+valor)
    _set_cell(row1_unique[0], "{{ pgto_entrada_valor }}")
    _set_cell(row1_unique[1], "{{ pgto_entrada_tipo }}")
    _set_cell(row1_unique[2], "{{ pgto_entrada_data }}")

    # Linha de parcelas (R2): unique cells são 0, 2, 4
    row2_unique = _unique_cells(t3_rows[2])
    _set_cell(row2_unique[0], "{{ pgto_modalidade }}")
    _set_cell(row2_unique[1], "{{ pgto_num_parcelas }}")
    _set_cell(row2_unique[2], "{{ pgto_data_primeira }}")

    # Grade de parcelas: linhas 3 a 10 (índices 3–10), 3 pares por linha
    # Cada linha tem 6 células independentes: (num, data, num, data, num, data)
    # Apenas as células de DATA (índices 1, 3, 5) recebem placeholder de data
    parcela_idx = 1
    for row_idx in range(3, 11):
        if parcela_idx > 24:
            break
        row = t3_rows[row_idx]
        cells = row.cells
        for col_data_idx in [1, 3, 5]:
            if parcela_idx > 24:
                break
            try:
                _set_cell(cells[col_data_idx], "{{ p%02d_data }}" % parcela_idx)
            except Exception:
                pass
            parcela_idx += 1

    # ── Parágrafos do corpo — data e assinatura ───────────────────────────────
    paras = doc.paragraphs

    # P122: "São José dos Campos - SP, 09 de junho de 2026."
    # → substituir data por placeholder
    for para in paras:
        t = para.text.strip()
        if t.startswith("São José dos Campos") and ("de 2026" in t or "de 20" in t):
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = "São José dos Campos - SP, {{ data_contrato }}."
            else:
                para.add_run("São José dos Campos - SP, {{ data_contrato }}.")
            break

    # P128: nome hardcoded do assinante da empresa (berê Ferreira Machado)
    # → substituir pelo nome do consultor/responsável
    for para in paras:
        t = para.text.strip()
        if "Ferreira Machado" in t or ("787.834" in t):
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = "{{ consultor_nome }}"
            else:
                para.add_run("{{ consultor_nome }}")
            break

    doc.save(DEST)
    print(f"Template salvo: {DEST}")

    # Verificação rápida
    doc2 = Document(DEST)
    p0_text = doc2.paragraphs[0].text
    assert "consultor_nome" in p0_text, f"Placeholder consultor_nome não encontrado: {p0_text}"
    t0_r1_c0 = doc2.tables[0].rows[1].cells[0].text
    assert "cliente_nome" in t0_r1_c0, f"Placeholder cliente_nome não encontrado: {t0_r1_c0}"
    t0_r1_c1 = doc2.tables[0].rows[1].cells[1].text
    assert "cliente_cpf" in t0_r1_c1, f"Placeholder cliente_cpf não encontrado: {t0_r1_c1}"
    t0_r2_c0 = doc2.tables[0].rows[2].cells[0].text
    assert "cliente_email" in t0_r2_c0, f"Placeholder cliente_email não encontrado: {t0_r2_c0}"
    t0_r2_c1 = doc2.tables[0].rows[2].cells[1].text
    assert "cliente_telefone" in t0_r2_c1, f"Placeholder cliente_telefone não encontrado: {t0_r2_c1}"
    print("Verificação OK — placeholders inseridos corretamente.")


if __name__ == "__main__":
    main()
