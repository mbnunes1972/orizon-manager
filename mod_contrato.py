"""
mod_contrato.py — Geração de contrato via template HTML/Markdown, renderizado em
PDF com WeasyPrint.

Ainda expõe um pequeno motor de substituição de marcadores em .docx
(_substituir_marcadores/_subst_paragrafo) e a conversão via LibreOffice
(_converter_pdf/_libreoffice_cmd/LibreOfficeIndisponivel) — não usados mais pelo
contrato, mas reaproveitados por mod_proposta.py (Proposta comercial, ainda em
.docx). Não remover sem migrar a Proposta primeiro.
"""

import os
import json
import platform
import subprocess
import hashlib
from datetime import datetime

_THIS_DIR            = os.path.dirname(os.path.abspath(__file__))
CONTRATOS_DIR        = os.path.join(_THIS_DIR, "CONTRATOS")
CONTRATO_TEMPLATE_DIR = os.path.join(_THIS_DIR, "contrato_template")

_TRACO = "--------"  # preenche slots de parcela inexistentes


# ── Utilitários ───────────────────────────────────────────────────────────────


def gerar_num_contrato(existing_nums, loja_codigo: str, data=None) -> str:
    """Próximo número de contrato no formato 'LOJA-AAAA-MM-DD-SEQ'.

    `existing_nums`: iterável com os num_contrato já existentes (qualquer loja).
    `loja_codigo`: código (3 letras) da loja — vem da tabela `lojas` (F3).
    A sequência (SEQ) é CONTÍNUA por código (máximo existente + 1).
    """
    data = data or datetime.now()
    cod = (loja_codigo or "").strip()
    pref = f"{cod}-"
    maxseq = 0
    for n in (existing_nums or []):
        if n and n.startswith(pref):
            tail = n.rsplit("-", 1)[-1]
            if tail.isdigit():
                maxseq = max(maxseq, int(tail))
    return f"{cod}-{data:%Y-%m-%d}-{maxseq + 1:03d}"


def calcular_hash_assinatura(nome: str, cpf: str, contrato_id: int, timestamp: str) -> str:
    dados = f"{nome}|{cpf}|{contrato_id}|{timestamp}"
    return hashlib.sha256(dados.encode("utf-8")).hexdigest()


def _formatar_valor(valor: float) -> str:
    return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _formatar_valor_str(v):
    """Aceita número ou string já formatada; devolve 'R$ x.xxx,xx' (ou '' se vazio)."""
    if v is None or v == "":
        return ""
    if isinstance(v, (int, float)):
        return _formatar_valor(v)
    return str(v).strip()


def ambientes_valor_contrato(itens, vavo, val_cont):
    """Distribui Val_Cont pelos ambientes, proporcional ao VAVA.

    itens: lista [(nome, VAVA_float), ...] na ordem do orçamento.
    Retorna [(nome, valor_float), ...] com Σ valor == round(val_cont, 2);
    o resíduo de arredondamento é absorvido pelo último ambiente.
    vavo<=0 → devolve os próprios VAVA arredondados (sem divisão por zero).
    """
    if not itens:
        return []
    alvo = round(float(val_cont or 0.0), 2)
    if not vavo or vavo <= 0:
        return [(n, round(float(v or 0.0), 2)) for n, v in itens]
    fator = alvo / vavo
    out = [(n, round(float(v or 0.0) * fator, 2)) for n, v in itens]
    resid = round(alvo - sum(v for _, v in out), 2)
    if resid:
        n_ult, v_ult = out[-1]
        out[-1] = (n_ult, round(v_ult + resid, 2))
    return out


def _formatar_data_br(data: str) -> str:
    """Converte ISO 'YYYY-MM-DD' → 'DD/MM/AAAA'. Datas já em DD/MM são passadas direto."""
    if not data:
        return "—"
    data = data.strip()
    # Já está no formato brasileiro
    if len(data) == 10 and data[2] == "/" and data[5] == "/":
        return data
    if len(data) >= 10:
        try:
            return datetime.strptime(data[:10], "%Y-%m-%d").strftime("%d/%m/%Y")
        except ValueError:
            pass
    return data or "—"


def _cel_amb(rotulo, valor):
    """Retorna as duas <td>s de um ambiente (rótulo e valor)."""
    return (f'<td class="amb-rotulo">{rotulo}</td>'
            f'<td class="amb-valor">{valor}</td>')


def _html_ambientes_linhas(itens_valores):
    """<tr>s da tabela de ambientes: 2 por linha; sobra ímpar em traços."""
    from html import escape
    n = len(itens_valores)
    n_linhas = max(1, (n + 1) // 2)
    linhas = []
    for k in range(n_linhas):
        cels = []
        for slot in (0, 1):
            idx = 2 * k + slot
            if idx < n:
                nome, val = itens_valores[idx]
                cels.append(_cel_amb(escape(str(nome)), _formatar_valor(val)))
            else:
                cels.append(_cel_amb(_TRACO, _TRACO))
        linhas.append("<tr>" + "".join(cels) + "</tr>")
    return "\n".join(linhas)


def _html_parcelas_linhas(pag):
    """<tr>s da grade de parcelas: 3 por linha, só linhas usadas, tracos no resto."""
    tipo = pag.get("tipo", "")
    num = pag.get("num_parcelas_int", 0)
    valores = pag.get("valores", [""] * 24)
    datas = pag.get("datas", [""] * 24)
    n_linhas = (num + 2) // 3  # ceil(num/3)
    linhas = []
    for gi in range(n_linhas):
        cels = []
        for j in range(3):
            p = gi * 3 + j + 1  # 1-based
            if p <= num and valores[p - 1]:
                val = valores[p - 1]
                data = "" if tipo == "cartao" else (datas[p - 1] or _TRACO)
            else:
                val, data = _TRACO, _TRACO
            cels.append(f'<td class="pc-valor">{val}</td>'
                        f'<td class="pc-data">{data}</td>')
        linhas.append("<tr>" + "".join(cels) + "</tr>")
    return "\n".join(linhas)


# ── Motor de substituição de marcadores [MARCADOR] ────────────────────────────

import re as _re_mark
_MARK_RE = _re_mark.compile(r'\[+\s*([A-Za-z0-9_ ]+?)\s*\]')

# ── Nível de cláusula + Markdown→HTML ──────────────────────────────────────────

import re as _re2

_RE_NUM = _re2.compile(r'^\s*(\d+(?:\.\d+)*)\.\s')
_RE_ALINEA = _re2.compile(r'^\s*[a-z]\)\s')


def _nivel_clausula(texto):
    """Detecta o nível de uma cláusula pela numeração literal.

    2. → 1, 2.3. → 2, 2.3.1. → 3, a) → 4, caso contrário None.
    """
    m = _RE_NUM.match(texto or "")
    if m:
        return m.group(1).count(".") + 1
    if _RE_ALINEA.match(texto or ""):
        return 4
    return None


def _inline_md(texto):
    """Formatação inline do Markdown (negrito/itálico), sem parsing de bloco.

    Protege runs de '_' (linhas de preenchimento, ex.: '(nome) __________') para o
    Markdown não os tratar como ênfase e destruí-los."""
    import markdown
    t = (texto or "").replace("_", "\x00u\x00")
    html = markdown.markdown(t).replace("\x00u\x00", "_")
    m = _re2.match(r'^<p>(.*)</p>$', html, _re2.DOTALL)
    return m.group(1) if m else html


def _html_corpo(md_texto):
    """Corpo (cláusulas com números literais) em HTML, linha a linha.
    Linhas '1.'/'1.1.'/'a)' viram <p class="cl-N"> (NÃO <li> de lista ordenada);
    '#'/'##' viram <h1>/<h2>; o número literal é preservado como texto e o
    Markdown é aplicado só ao texto após o número (inline)."""
    linhas = []
    for bruta in md_texto.splitlines():
        t = bruta.strip()
        if not t:
            continue
        mh = _re2.match(r'^(#{1,6})\s+(.*)$', t)
        if mh:
            lvl = len(mh.group(1))
            linhas.append(f"<h{lvl}>{_inline_md(mh.group(2))}</h{lvl}>")
            continue
        if t.startswith("E assim") and "firmam as PARTES" in t:
            # parágrafo de fecho: espaço acima via CSS
            linhas.append(f'<p class="fecho">{_inline_md(t)}</p>')
            continue
        if "[DATA_CONTRATO]" in t:
            # local/data do fecho (área de assinaturas): espaço acima via CSS
            linhas.append(f'<p class="data-fecho">{_inline_md(t)}</p>')
            continue
        if _re2.match(r'^_{5,}$', t):
            # linha de assinatura (traço para assinar): espaço acima via CSS
            linhas.append(f'<p class="assinatura">{t}</p>')
            continue
        alinea = _RE_ALINEA.match(t)
        nivel = _nivel_clausula(t)
        if alinea or nivel is not None:
            if alinea:                       # a), b)… -> alínea
                mnum, classe = alinea, "cl-alinea"
            else:                            # cláusula numérica (nível ≥3 usa cl-3)
                mnum, classe = _RE_NUM.match(t), f"cl-{min(nivel, 3)}"
            prefixo = t[:mnum.end()]         # ex.: "1.1. " ou "a) "
            resto = t[mnum.end():]
            linhas.append(f'<p class="{classe}">{prefixo}{_inline_md(resto)}</p>')
        else:
            linhas.append(f'<p>{_inline_md(t)}</p>')
    return "\n".join(linhas)


def _aplica_mark(texto, mapping):
    def repl(m):
        chave = m.group(1).strip().upper().replace(" ", "_")
        return mapping[chave] if chave in mapping else m.group(0)
    return _MARK_RE.sub(repl, texto)


def _substituir_marcadores_html(html, mapping):
    """Substitui [MARCADOR] (case-insensitive, tolera '[[') numa string HTML/texto.
    Chaves do mapping em MAIÚSCULAS sem colchetes. Marcador sem chave é mantido."""
    def repl(m):
        chave = m.group(1).strip().upper().replace(" ", "_")
        return mapping[chave] if chave in mapping else m.group(0)
    return _MARK_RE.sub(repl, html)


# ── Motor de substituição de marcadores em .docx (legado — usado por mod_proposta) ──

def _subst_paragrafo(par, mapping, coletor=None):
    """Reconstrói o parágrafo em segmentos (texto fixo × valor substituído),
    preservando a formatação (``rPr``) de CADA run original.

    Cada trecho de texto herda o estilo do run de onde veio (rótulo cinza pequeno
    continua cinza pequeno; valor em negrito continua negrito). O valor substituído
    herda o estilo do run onde o marcador começava. Quando ``coletor`` é fornecido,
    o run do VALOR substituído é anexado a ele (isola o valor para torná-lo editável).
    O texto resultante (``par.text``) é idêntico ao da substituição direta.
    """
    txt = par.text
    if "[" not in txt:
        return
    import copy as _copy
    from docx.oxml.ns import qn as _qn

    runs = list(par.runs)
    # mapa caractere -> run original de onde veio (para preservar formatação por run)
    char_run = []
    for r in runs:
        char_run.extend([r] * len(r.text))
    if len(char_run) != len(txt):
        # desalinhamento (conteúdo não-run: hyperlink/campo) → não arrisca reformatar
        char_run = None

    def _src(i):
        if char_run and 0 <= i < len(char_run):
            return char_run[i]
        return runs[0] if runs else None

    # segmentos: (texto, run_fonte, eh_valor) — quebrando trechos fixos por run de origem
    segs = []
    def _emit_fixo(sub, ini):
        if not sub:
            return
        if not char_run:
            segs.append((sub, runs[0] if runs else None, False))
            return
        k = 0
        for j in range(1, len(sub) + 1):
            if j == len(sub) or char_run[ini + j] is not char_run[ini + j - 1]:
                segs.append((sub[k:j], char_run[ini + k], False))
                k = j

    pos = 0
    for m in _MARK_RE.finditer(txt):
        if m.start() > pos:
            _emit_fixo(txt[pos:m.start()], pos)
        chave = m.group(1).strip().upper().replace(" ", "_")
        if chave in mapping:
            segs.append((mapping[chave], _src(m.start()), True))
        else:
            _emit_fixo(m.group(0), m.start())
        pos = m.end()
    if pos < len(txt):
        _emit_fixo(txt[pos:], pos)

    # captura o rPr de cada run fonte ANTES de remover os runs originais
    rpr_por_run = {}
    for _, src, _v in segs:
        if src is not None and id(src) not in rpr_por_run:
            el = src._r.find(_qn('w:rPr'))
            rpr_por_run[id(src)] = _copy.deepcopy(el) if el is not None else None

    # limpa runs existentes e recria os segmentos preservando a formatação de origem
    for r in runs:
        r._r.getparent().remove(r._r)
    for seg_txt, src, eh_valor in segs:
        run = par.add_run(seg_txt)
        rpr = rpr_por_run.get(id(src)) if src is not None else None
        if rpr is not None:
            velho = run._r.find(_qn('w:rPr'))
            if velho is not None:
                run._r.remove(velho)
            run._r.insert(0, _copy.deepcopy(rpr))
        if eh_valor and coletor is not None:
            coletor.append(run)


def _substituir_marcadores(doc, mapping, coletor=None):
    """Substitui [MARCADOR] (case-insensitive, tolera '[[') no corpo, tabelas e headers.
    Chaves do mapping SEM colchetes, em MAIÚSCULAS. Marcador sem chave é mantido."""
    for par in doc.paragraphs:
        _subst_paragrafo(par, mapping, coletor)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    _subst_paragrafo(par, mapping, coletor)
    from docx.text.paragraph import Paragraph as _Paragraph
    _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for sec in doc.sections:
        for hdr in (sec.header, sec.first_page_header, sec.even_page_header):
            for par in hdr.paragraphs:
                _subst_paragrafo(par, mapping)
            for tbl in hdr.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for par in cell.paragraphs:
                            _subst_paragrafo(par, mapping)
            # Text boxes (wps:txbx / mc:AlternateContent) hold markers in w:txbxContent
            for txbx in hdr._element.findall(f'.//{{{_W}}}txbxContent'):
                for p_el in txbx.findall(f'{{{_W}}}p'):
                    _subst_paragrafo(_Paragraph(p_el, None), mapping)


# ── Parser de pagamento ───────────────────────────────────────────────────────

_FORMA_LABELS = {
    "pix": "Pix",
    "ted": "TED",
    "transferencia": "TED",
    "boleto": "Boleto",
    "cheque": "Cheque",
    "dinheiro": "Dinheiro",
    "cartao_credito": "Cartão de Crédito",
    "cartao_debito": "Cartão de Débito",
    "debito_automatico": "Débito Automático",
}


def _forma_label(codigo: str) -> str:
    """Converte código de forma de pagamento em rótulo pt-BR. Idempotente:
    um rótulo já formatado (não encontrado no mapa) é devolvido como veio."""
    if not codigo:
        return ""
    return _FORMA_LABELS.get(str(codigo).strip().lower(), str(codigo).strip())


def _parse_pagamento(pag_json_str: str) -> dict:
    """
    Normaliza o JSON de pagamento capturado pelo frontend (_capturarPagamento).

    Formato de entrada (do frontend):
      { tipo, nome_forma, entrada_valor, entrada_data, entrada_forma,
        parcelas: [{seq, descricao, data, valor, forma}],
        texto (apenas cartão) }

    Retorna dict normalizado com campos prontos para preencher a capa.
    """
    try:
        pag = json.loads(pag_json_str) if pag_json_str else {}
    except Exception:
        pag = {}

    tipo         = (pag.get("tipo") or "").lower()
    nome_forma   = pag.get("nome_forma") or ""
    entrada_val  = float(pag.get("entrada_valor") or 0)
    entrada_data = _formatar_data_br(pag.get("entrada_data") or "")
    entrada_tipo = _forma_label(pag.get("entrada_forma") or pag.get("entrada_tipo") or "")
    parcelas     = pag.get("parcelas") or []
    num_parcelas = len(parcelas)

    if tipo == "cartao":
        forma_parcela = "Cartão de Crédito"
    elif parcelas:
        forma_parcela = _forma_label(parcelas[0].get("forma") or "")
    else:
        forma_parcela = ""

    # Grade p01..p24 — datas e valores diretamente das parcelas reais capturadas
    datas, valores = [], []
    for p in parcelas:
        if tipo == "cartao":
            datas.append("")   # cartão: parcelas não têm data na grade
        else:
            datas.append(_formatar_data_br(p.get("data") or ""))
        valores.append(_formatar_valor_str(p.get("valor")))
    datas   = (datas   + [""] * 24)[:24]
    valores = (valores + [""] * 24)[:24]

    if tipo == "cartao" and num_parcelas == 1:
        num_parcelas_disp = "à vista"
    elif num_parcelas:
        num_parcelas_disp = str(num_parcelas)
    else:
        num_parcelas_disp = "—"

    total_cliente = pag.get("total_cliente") or 0
    return {
        "tipo":             tipo,
        "nome_forma":       nome_forma,
        "entrada_valor":    _formatar_valor(entrada_val),
        "entrada_tipo":     entrada_tipo,
        "entrada_data":     entrada_data,
        "modalidade":       nome_forma,
        "num_parcelas":     num_parcelas_disp,
        "num_parcelas_int": num_parcelas,
        "data_primeira":    (datas[0] if datas and datas[0] else ""),
        "datas":            datas,          # lista de 24 strings (data ou "")
        "valores":          valores,        # lista de 24 strings (valor ou "")
        "valor_contrato":   _formatar_valor(total_cliente),
        "texto_cartao":     pag.get("texto_cartao") or "",
        "forma_parcela":    forma_parcela,
    }


# ── Guard de desatualização ───────────────────────────────────────────────────

def contrato_desatualizado(pagamento_snapshot_json, forma_pagamento_atual_json):
    """True se o pagamento mudou após a geração do contrato (modalidade 'tipo' ou 'total_cliente').
    Compara 'tipo' e 'total_cliente' (arredondado a 2 casas). Retorna False se não der pra comparar."""
    import json as _json
    def _p(x):
        if not x: return None
        try: return _json.loads(x) if isinstance(x, str) else x
        except Exception: return None
    a = _p(pagamento_snapshot_json); b = _p(forma_pagamento_atual_json)
    if not a or not b:
        return False
    if (a.get("tipo") or "") != (b.get("tipo") or ""):
        return True
    return round(float(a.get("total_cliente") or 0), 2) != round(float(b.get("total_cliente") or 0), 2)


# ── Validação de dados do cliente ─────────────────────────────────────────────

def validar_cliente_para_contrato(cliente: dict) -> list:
    """
    Retorna a lista de rótulos dos campos obrigatórios que estão vazios para
    gerar um contrato a partir do dict do cliente (formato _cliente_dict).

    Lista vazia → cliente está completo e pode gerar o contrato.

    Regra: identificação + endereço residencial são sempre obrigatórios.
    O endereço de instalação só é cobrado quando NÃO for o mesmo do residencial
    (inst_mesmo_residencial falso) — quando é o mesmo, o residencial é reutilizado.
    Complemento é opcional em ambos.
    """
    obrigatorios = [
        ("nome",       "Nome"),
        ("cpf",        "CPF"),
        ("email",      "E-mail"),
        ("telefone",   "Telefone"),
        ("logradouro", "Logradouro (residencial)"),
        ("numero",     "Número (residencial)"),
        ("bairro",     "Bairro (residencial)"),
        ("cidade",     "Cidade (residencial)"),
        ("cep",        "CEP (residencial)"),
        ("estado",     "Estado/UF (residencial)"),
    ]

    inst_mesmo = cliente.get("inst_mesmo_residencial", True)
    if not inst_mesmo:
        obrigatorios += [
            ("inst_logradouro", "Logradouro (instalação)"),
            ("inst_numero",     "Número (instalação)"),
            ("inst_bairro",     "Bairro (instalação)"),
            ("inst_cidade",     "Cidade (instalação)"),
            ("inst_cep",        "CEP (instalação)"),
            ("inst_uf",         "UF (instalação)"),
        ]

    faltando = []
    for campo, rotulo in obrigatorios:
        valor = cliente.get(campo)
        if not (valor and str(valor).strip()):
            faltando.append(rotulo)
    return faltando


# ── Validação de dados da loja ────────────────────────────────────────────────

_TEM_DIGITO = _re_mark.compile(r"\d")   # CPF real tem ao menos um dígito


def validar_loja_para_contrato(loja: dict) -> list:
    """Rótulos dos campos obrigatórios da loja que estão vazios para gerar o contrato.

    Lista vazia → loja completa. O CPF de testemunha sem nenhum dígito
    (placeholder 'xxx.xxx.xxx-xx') conta como faltando. `complemento` é opcional.
    """
    loja = loja or {}
    obrigatorios = [
        ("nome",             "Nome da empresa"),
        ("cnpj",             "CNPJ"),
        ("codigo",           "Código da loja"),
        ("telefone",         "Telefone"),
        ("email",            "E-mail"),
        ("cep",              "CEP"),
        ("logradouro",       "Logradouro"),
        ("numero",           "Número"),
        ("bairro",           "Bairro"),
        ("cidade",           "Cidade"),
        ("estado",           "Estado/UF"),
        ("testemunha1_nome", "Nome da Testemunha 1"),
        ("testemunha2_nome", "Nome da Testemunha 2"),
    ]
    faltando = []
    for campo, rotulo in obrigatorios:
        v = loja.get(campo)
        if not (v and str(v).strip()):
            faltando.append(rotulo)
    for campo, rotulo in [("testemunha1_cpf", "CPF da Testemunha 1"),
                          ("testemunha2_cpf", "CPF da Testemunha 2")]:
        v = (loja.get(campo) or "").strip()
        if not _TEM_DIGITO.search(v):
            faltando.append(rotulo)
    return faltando


# ── Preenchimento dinâmico do modelo ─────────────────────────────────────────

def _montar_mapping(ctx, pag):
    """Monta o dicionário {MARCADOR: valor} para _substituir_marcadores_html().

    Chaves em MAIÚSCULAS sem colchetes, casando com os marcadores do template
    HTML do contrato. A grade de parcelas e os ambientes são montados à parte
    (_html_parcelas_linhas/_html_ambientes_linhas); aqui só entram os campos de
    cabeçalho/identificação. Os dados da loja vêm de ctx['loja'] (F3).
    """
    loja = ctx.get("loja") or {}
    t1n = loja.get("testemunha1_nome", "") or ""
    t1c = loja.get("testemunha1_cpf", "") or ""
    t2n = loja.get("testemunha2_nome", "") or ""
    t2c = loja.get("testemunha2_cpf", "") or ""
    return {
        "NUM_CONTRATO":     str(ctx.get("num_contrato", "") or ""),
        "DATA_CONTRATO":    str(ctx.get("data_contrato", "") or ""),
        "NOME_CLIENTE":     ctx.get("cliente_nome", "") or "",
        "CPF":              ctx.get("cliente_cpf", "") or "",
        "EMAIL":            ctx.get("cliente_email", "") or "",
        "TELEFONE":         ctx.get("cliente_telefone", "") or "",
        "RES_LOGRADOURO":   ctx.get("res_logradouro", "") or "",
        "RES_NUMERO":       ctx.get("res_numero", "") or "",
        "RES_COMPLEMENTO":  ctx.get("res_complemento", "") or "",
        "RES_BAIRRO":       ctx.get("res_bairro", "") or "",
        "RES_CIDADE":       ctx.get("res_cidade", "") or "",
        "RES_CEP":          ctx.get("res_cep", "") or "",
        "RES_UF":           ctx.get("res_uf", "") or "",
        "INST_LOGRADOURO":  ctx.get("inst_logradouro", "") or "",
        "INST_NUMERO":      ctx.get("inst_numero", "") or "",
        "INST_COMPLEMENTO": ctx.get("inst_complemento", "") or "",
        "INST_BAIRRO":      ctx.get("inst_bairro", "") or "",
        "INST_CIDADE":      ctx.get("inst_cidade", "") or "",
        "INST_CEP":         ctx.get("inst_cep", "") or "",
        "INST_UF":          ctx.get("inst_uf", "") or "",
        "VALOR_ENTRADA":    pag.get("entrada_valor", "") or "",
        "FORMA_ENTRADA":    pag.get("entrada_tipo", "") or "",
        "DATA_ENTRADA":     pag.get("entrada_data", "") or "",
        "MODALIDADE":       pag.get("nome_forma", "") or "",
        "NUM_PARCELAS":     pag.get("num_parcelas", "") or "",
        "TIPO":             pag.get("forma_parcela", "") or "",
        "TOTAL_CONTRATO":   pag.get("valor_contrato", "") or "",
        "CONSULTOR_NOME":     ctx.get("consultor_nome", "") or "",
        "CONSULTOR_TELEFONE": ctx.get("consultor_tel", "") or "",
        # Identificador da rede sob o logo (parâmetro de config do contrato;
        # fallback p/ a cidade da loja enquanto o campo de config não é definido)
        "REDE_IDENTIFICADOR": ctx.get("rede_identificador", "") or loja.get("cidade", "") or "",
        "TESTEMUNHA_1_NOME": t1n,
        "TESTEMUNHA_1_DOC":  t1c,
        "TESTEMUNHA_2_NOME": t2n,
        "TESTEMUNHA_2_DOC":  t2c,
        "NOME_TESTEMUNHA_1": t1n,
        "NOME_TESTEMUNHA2":  t2n,
        "NOME_TESTEMUNHA_2": t2n,
        "NOME_EMPRESA":      loja.get("nome", "") or "",
        "CNPJ_EMPRESA":      loja.get("cnpj", "") or "",
        "CPF_CLIENTE":       ctx.get("cliente_cpf", "") or "",
        "CPF_TESTEMUNHA_1":  t1c,
        "CPF_TESTEMUNHA_2":  t2c,
    }


def _html_capa(ctx):
    """Monta o HTML da capa (identificação, endereços, ambientes e pagamento).

    Consome _html_ambientes_linhas/_html_parcelas_linhas para as linhas dinâmicas
    e deixa os demais campos como [MARCADORES] a serem resolvidos por
    _substituir_marcadores_html() com o mapping de _montar_mapping().
    """
    amb = _html_ambientes_linhas(ctx.get("_ambientes") or [])
    parc = _html_parcelas_linhas(ctx.get("_pag") or {})
    return f"""
<div id="cabecalho">
  <div class="cab-esq">
    <img class="logo" src="logo_dalmobile.png">
    <div class="cab-rede">[REDE_IDENTIFICADOR]</div>
  </div>
  <div class="cab-dir">[NUM_CONTRATO]<br>[DATA_CONTRATO]</div>
</div>

<div class="consultor">Consultor: [CONSULTOR_NOME] &nbsp;&nbsp;&nbsp; Telefone: [CONSULTOR_TELEFONE]</div>

<div class="secao">
  <div class="titulo">1. Identificação do Cliente</div>
  <table><tr>
    <td><span class="rotulo">Nome</span><span class="valor">[NOME_CLIENTE]</span></td>
    <td><span class="rotulo">CPF/CNPJ</span><span class="valor">[CPF]</span></td>
  </tr><tr>
    <td><span class="rotulo">E-mail</span><span class="valor">[EMAIL]</span></td>
    <td><span class="rotulo">Telefone</span><span class="valor">[TELEFONE]</span></td>
  </tr></table>
</div>

<div class="secao">
  <div class="titulo">2. Endereço Residencial</div>
  <table><tr>
    <td colspan="3"><span class="rotulo">Logradouro</span><span class="valor">[RES_LOGRADOURO]</span></td>
  </tr><tr>
    <td><span class="rotulo">Número</span><span class="valor">[RES_NUMERO]</span></td>
    <td><span class="rotulo">Complemento</span><span class="valor">[RES_COMPLEMENTO]</span></td>
    <td><span class="rotulo">Bairro</span><span class="valor">[RES_BAIRRO]</span></td>
  </tr><tr>
    <td><span class="rotulo">Cidade</span><span class="valor">[RES_CIDADE]</span></td>
    <td><span class="rotulo">CEP</span><span class="valor">[RES_CEP]</span></td>
    <td><span class="rotulo">Estado/UF</span><span class="valor">[RES_UF]</span></td>
  </tr></table>
</div>

<div class="secao">
  <div class="titulo">3. Endereço de Instalação</div>
  <table><tr>
    <td colspan="3"><span class="rotulo">Logradouro</span><span class="valor">[INST_LOGRADOURO]</span></td>
  </tr><tr>
    <td><span class="rotulo">Número</span><span class="valor">[INST_NUMERO]</span></td>
    <td><span class="rotulo">Complemento</span><span class="valor">[INST_COMPLEMENTO]</span></td>
    <td><span class="rotulo">Bairro</span><span class="valor">[INST_BAIRRO]</span></td>
  </tr><tr>
    <td><span class="rotulo">Cidade</span><span class="valor">[INST_CIDADE]</span></td>
    <td><span class="rotulo">CEP</span><span class="valor">[INST_CEP]</span></td>
    <td><span class="rotulo">Estado/UF</span><span class="valor">[INST_UF]</span></td>
  </tr></table>
</div>

<div class="secao">
  <div class="titulo">4. Ambientes do Projeto</div>
  <table>
    {amb}
    <tr><td class="total-lbl" colspan="2">VALOR DO CONTRATO</td>
        <td class="total-val" colspan="2">[TOTAL_CONTRATO]</td></tr>
  </table>
</div>

<div class="secao">
  <div class="titulo">5. Forma de Pagamento</div>
  <table><tr>
    <td><span class="rotulo">Entrada</span><span class="valor">[VALOR_ENTRADA]</span></td>
    <td><span class="rotulo">Tipo</span><span class="valor">[FORMA_ENTRADA]</span></td>
    <td><span class="rotulo">Data</span><span class="valor">[DATA_ENTRADA]</span></td>
  </tr><tr>
    <td><span class="rotulo">Modalidade</span><span class="valor">[MODALIDADE]</span></td>
    <td><span class="rotulo">Parcelas</span><span class="valor">[NUM_PARCELAS]</span></td>
    <td><span class="rotulo">Valor do Contrato</span><span class="valor">[TOTAL_CONTRATO]</span></td>
  </tr></table>
  <table>{parc}</table>
</div>

<div class="quebra-capa"></div>
"""


def _carregar_md():
    """Lê contrato_template/contrato.md; retorna "" se o arquivo ainda não existir."""
    p = os.path.join(CONTRATO_TEMPLATE_DIR, "contrato.md")
    return open(p, encoding="utf-8").read() if os.path.exists(p) else ""


def _montar_html_contrato(ctx):
    """Monta o HTML final do contrato: capa + corpo, com [MARCADORES] substituídos."""
    from html import escape
    pag = ctx.get("_pag", {})
    mapping = _montar_mapping(ctx, pag)
    mapping["TEXTO_COMPLEMENTAR"] = ctx.get("adendo", "") or ""
    # escapa os VALORES dos marcadores (dados do cliente/loja/adendo) para HTML —
    # nomes/endereços/adendo (texto livre) não podem corromper o documento.
    # (o escape é aqui, no caminho HTML; _montar_mapping é compartilhado com a
    #  proposta, que ainda gera .docx e não deve receber texto escapado.)
    mapping = {k: escape(str(v)) for k, v in mapping.items()}
    if mapping["TEXTO_COMPLEMENTAR"]:
        mapping["TEXTO_COMPLEMENTAR"] = mapping["TEXTO_COMPLEMENTAR"].replace("\n", "<br>")
    shell = open(os.path.join(CONTRATO_TEMPLATE_DIR, "contrato.html"),
                 encoding="utf-8").read()
    capa = _html_capa(ctx)
    corpo = _html_corpo(_carregar_md())
    html = shell.replace("<!--CAPA-->", capa).replace("<!--CORPO-->", corpo)
    return _substituir_marcadores_html(html, mapping)


def construir_contexto(cliente: dict, usuario: dict, forma_pagamento_json: str, loja: dict = None) -> dict:
    """Monta o dicionário completo para preencher o contrato.

    `loja`: dict com dados da loja (F3) — usado como fallback de telefone/email do
    consultor e injetado em ctx['loja'] para _montar_mapping(). Opcional; se None,
    os campos de empresa/testemunhas ficarão em branco no contrato.
    """
    loja = loja or {}
    inst_mesmo = cliente.get("inst_mesmo_residencial", True)
    if inst_mesmo:
        inst = {k: cliente.get(r, "") for k, r in [
            ("logradouro","logradouro"), ("numero","numero"),
            ("complemento","complemento"), ("bairro","bairro"),
            ("cidade","cidade"), ("cep","cep"), ("uf","estado"),
        ]}
    else:
        inst = {
            "logradouro":  cliente.get("inst_logradouro",  ""),
            "numero":      cliente.get("inst_numero",      ""),
            "complemento": cliente.get("inst_complemento", ""),
            "bairro":      cliente.get("inst_bairro",      ""),
            "cidade":      cliente.get("inst_cidade",      ""),
            "cep":         cliente.get("inst_cep",         ""),
            "uf":          cliente.get("inst_uf",          ""),
        }

    pag = _parse_pagamento(forma_pagamento_json)
    datas = pag["datas"]

    ctx = {
        "consultor_nome":  usuario.get("nome",     "") or "",
        "consultor_tel":   (usuario.get("telefone") or "").strip() or (loja.get("telefone") or ""),
        "consultor_email": (usuario.get("email")    or "").strip() or (loja.get("email")    or ""),
        "cliente_nome":    cliente.get("nome",     "") or "",
        "cliente_cpf":     cliente.get("cpf",      "") or "",
        "cliente_email":   cliente.get("email",    "") or "",
        "cliente_telefone":cliente.get("telefone", "") or "",
        "res_logradouro":  cliente.get("logradouro",  "") or "",
        "res_numero":      cliente.get("numero",      "") or "",
        "res_complemento": cliente.get("complemento", "") or "",
        "res_bairro":      cliente.get("bairro",      "") or "",
        "res_cidade":      cliente.get("cidade",      "") or "",
        "res_cep":         cliente.get("cep",         "") or "",
        "res_uf":          cliente.get("estado",      "") or "",
        "inst_logradouro":  inst["logradouro"],
        "inst_numero":      inst["numero"],
        "inst_complemento": inst["complemento"],
        "inst_bairro":      inst["bairro"],
        "inst_cidade":      inst["cidade"],
        "inst_cep":         inst["cep"],
        "inst_uf":          inst["uf"],
        # Pagamento — chaves planas para compatibilidade com templates antigos
        "pgto_entrada_valor":  pag["entrada_valor"],
        "pgto_entrada_tipo":   pag["entrada_tipo"],
        "pgto_entrada_data":   pag["entrada_data"],
        "pgto_modalidade":     pag["modalidade"],
        "pgto_num_parcelas":   pag["num_parcelas"],
        "pgto_data_primeira":  pag["data_primeira"],
        "data_contrato":       datetime.now().strftime("%d/%m/%Y"),
        "_pag":                pag,   # dict normalizado para gerar_pdf_contrato()
        "loja":                loja,
    }
    # Grade de datas p01..p24
    for i in range(24):
        ctx[f"p{i+1:02d}_data"] = datas[i]
    return ctx


def gerar_pdf_contrato(contrato_id: int, ctx: dict, destino: str = None) -> str:
    """Renderiza o contrato (HTML -> PDF) via WeasyPrint. Retorna o caminho do PDF."""
    from weasyprint import HTML
    destino = destino or CONTRATOS_DIR
    os.makedirs(destino, exist_ok=True)
    html = _montar_html_contrato(ctx)
    pdf_path = os.path.join(destino, f"contrato_{contrato_id}.pdf")
    HTML(string=html, base_url=CONTRATO_TEMPLATE_DIR).write_pdf(pdf_path)
    return pdf_path


# ── LibreOffice (legado — usado por mod_proposta para converter proposta.docx) ─

class LibreOfficeIndisponivel(Exception):
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        super().__init__(
            "LibreOffice não encontrado no servidor.\n"
            "O arquivo Word (.docx) foi gerado e está disponível para download.\n"
            "Instale o LibreOffice no servidor para gerar PDF automaticamente."
        )


def _libreoffice_cmd() -> str:
    if platform.system() == "Windows":
        for p in [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]:
            if os.path.exists(p):
                return p
    return "libreoffice"


def _converter_pdf(docx_path: str, outdir: str = None) -> str:
    """Converte um .docx EXISTENTE em PDF (não regenera). Retorna o caminho do PDF.
    outdir default = CONTRATOS_DIR (comportamento do contrato inalterado)."""
    destino = outdir or CONTRATOS_DIR
    try:
        subprocess.run(
            [_libreoffice_cmd(), "--headless", "--convert-to", "pdf",
             "--outdir", destino, docx_path],
            check=True, capture_output=True, timeout=120,
        )
    except FileNotFoundError:
        raise LibreOfficeIndisponivel(docx_path)
    except subprocess.CalledProcessError as e:
        raise RuntimeError(
            f"LibreOffice falhou:\n{e.stderr.decode(errors='replace')}"
        ) from e
    except subprocess.TimeoutExpired:
        raise RuntimeError("LibreOffice demorou mais de 120s")

    base = os.path.splitext(os.path.basename(docx_path))[0]
    return os.path.join(destino, f"{base}.pdf")
