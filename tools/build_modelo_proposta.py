"""Gera modelo_proposta.docx (template da Proposta comercial) com marcadores [MARCADOR].
Baseado na 1a pagina do contrato: cabecalho da loja, partes, oferta (ambientes/valores),
condicoes de pagamento, validade. Rode: python3 tools/build_modelo_proposta.py"""
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "modelo_proposta.docx")

doc = Document()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("PROPOSTA COMERCIAL"); r.bold = True; r.font.size = None

doc.add_paragraph("[NOME_EMPRESA] — CNPJ [CNPJ_EMPRESA]")
doc.add_paragraph("")
doc.add_paragraph("Cliente: [NOME_CLIENTE]    CPF: [CPF_CLIENTE]")
doc.add_paragraph("E-mail: [EMAIL]    Telefone: [TELEFONE]")
doc.add_paragraph("Consultor: [CONSULTOR_NOME]")
doc.add_paragraph("")
doc.add_paragraph("Ambientes:")
doc.add_paragraph("[AMBIENTES_LISTA]")
doc.add_paragraph("")
doc.add_paragraph("Valor bruto: [VALOR_BRUTO]")
doc.add_paragraph("Desconto: [DESCONTO_PCT]")
doc.add_paragraph("Valor total: [VALOR_TOTAL]")
doc.add_paragraph("")
doc.add_paragraph("Condições de pagamento: [MODALIDADE]")
doc.add_paragraph("Entrada: [VALOR_ENTRADA] ([FORMA_ENTRADA]) em [DATA_ENTRADA]")
doc.add_paragraph("Parcelas: [NUM_PARCELAS]")
doc.add_paragraph("")
doc.add_paragraph("[VALIDADE]")

doc.save(OUT)
print("gerado:", OUT)
